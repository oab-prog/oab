import os
import json
import time
import datetime
import threading
import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext, ttk
import docx
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from google import genai
from google.genai import types
import re # O motor do Escudo Anti-Alucinação

API_KEY_LOCAL = os.getenv("GEMINI_API_KEY") 

# ==========================================
# O ESCUDO DE CÓDIGO (REGEX) - A PORTA DE AÇO
# ==========================================
def aplicar_escudo_no_json(dados):
    # Caça padrões inventados de jurisprudência Trabalhista (RO, RR, AIRR, TRT, TST)
    padrao_jurisprudencia = r'(?i)\b(?:HC|REsp|RHC|AgRg|AREsp|Apelação|Agravo|RO|RR|AIRR|Processo n[º°]?\s*)\s*[\d\.\-\/]+\/[A-Z]{2}\b'
    
    if isinstance(dados, dict):
        return {k: aplicar_escudo_no_json(v) for k, v in dados.items()}
    elif isinstance(dados, list):
        return [aplicar_escudo_no_json(i) for i in dados]
    elif isinstance(dados, str):
        return re.sub(padrao_jurisprudencia, "[AVISO: NUMERACAO OMITIDA - BUSQUE NA INTEGRA NO JUSBRASIL/TST]", dados)
    else:
        return dados

def analisar_lote_ia(caminhos_txt, fatos, callback_progresso, callback_sucesso, callback_erro):
    try:
        client = genai.Client(api_key=API_KEY_LOCAL.strip())
        filtros_seguranca = [
            types.SafetySetting(category="HARM_CATEGORY_HATE_SPEECH", threshold="BLOCK_NONE"),
            types.SafetySetting(category="HARM_CATEGORY_HARASSMENT", threshold="BLOCK_NONE"),
            types.SafetySetting(category="HARM_CATEGORY_SEXUALLY_EXPLICIT", threshold="BLOCK_NONE"),
            types.SafetySetting(category="HARM_CATEGORY_DANGEROUS_CONTENT", threshold="BLOCK_NONE"),
        ]

        resumos_volumes = []
        total_vols = len(caminhos_txt)

        # FASE 1: DETETIVES TRABALHISTAS
        for idx, caminho in enumerate(caminhos_txt):
            vol_num = idx + 1
            callback_progresso(vol_num, total_vols, f"Fase 1: Auditando Provas e Cartões de Ponto no Vol {vol_num} de {total_vols}...")
            
            gemini_file = client.files.upload(file=caminho, config={'mime_type': 'text/plain'})
            while True:
                f_info = client.files.get(name=gemini_file.name)
                if "FAILED" in str(f_info.state).upper(): raise Exception(f"Falha no Volume {vol_num}.")
                if "ACTIVE" in str(f_info.state).upper(): break
                time.sleep(2)
            
            conteudo_txt = types.Part.from_uri(file_uri=f_info.uri, mime_type='text/plain')
            
            prompt_detetive = f"DIRECIONAMENTO DA ESTRATÉGIA:\n{fatos}\n\nLeia este volume Trabalhista. Extraia informações cruciais sobre: jornada de trabalho, horas extras, verbas rescisórias, insalubridade/periculosidade, depoimentos e contradições documentais. É OBRIGATÓRIO citar as fls. de cada achado. Seja robótico e exato. Se a informação não existir nos autos, responda: INFORMAÇÃO NÃO ENCONTRADA. NÃO INVENTE DADOS."
            config_detetive = types.GenerateContentConfig(temperature=0.0, safety_settings=filtros_seguranca)
            response_detetive = client.models.generate_content(model='gemini-2.5-flash', contents=[conteudo_txt, prompt_detetive], config=config_detetive)
            
            resumos_volumes.append(f"--- ACHADOS TRABALHISTAS DO VOLUME {vol_num} ---\n{response_detetive.text.strip()}\n")
            client.files.delete(name=gemini_file.name)

        # FASE 2: MESTRE JSON TRABALHISTA
        callback_progresso(total_vols, total_vols, "Fase 2: Estruturando Dossiê Jurídico Trabalhista...")
        texto_consolidado = "\n".join(resumos_volumes)

        instrucao_sistema_mestre = """
        Você é o M.A | JUS IA EXPERIENCE, um Auditor Forense Especialista em Direito do Trabalho (CLT).
        
        >>> TRAVA ABSOLUTA ANTI-ALUCINAÇÃO (TOLERÂNCIA ZERO) <<<
        1. PROIBIDO inventar fatos, salários, horários, locais ou testemunhos que não estejam EXPLICITAMENTE ESCRITOS nos dados.
        2. VALIDAR APENAS legislação trabalhista (CLT e Súmulas/OJs do TST).
        3. Se a informação não existir nos autos, responda: INFORMAÇÃO NÃO ENCONTRADA.
        4. OBRIGATÓRIO citar as fls. de cada achado.

        >>> RIGOR ESTATÍSTICO OAB (PESO DE CORREÇÃO) <<<
        Foque com RIGOR na estrutura de: Recurso Ordinário (Rigor 13), Contestação (Rigor 13), Reclamatória (Rigor 8).
        
        >>> REGRA DE FORMATAÇÃO JSON <<<
        O retorno DEVE ser EXCLUSIVAMENTE um objeto JSON válido.
        USE ESTA ESTRUTURA:
        {
          "resumo_estrategico": "Diagnóstico do processo trabalhista citando fls.",
          "analise_foco_cirurgico": ["Ponto trabalhista 1 com fls.", "Ponto 2 com fls."],
          "jurimetria": "Análise do juízo/vara do trabalho.",
          "resumo_cliente": "E-mail cordial para o cliente (Reclamante ou Reclamada).",
          "timeline": [{"data": "DD/MM/AAAA", "evento": "Descrição (Ex: Admissão, Demissão) e fls."}],
          "vulnerabilidades_contraparte": ["Falhas em cartões de ponto, nulidades ou prescrição com fls."],
          "checklist": ["Providência (Ex: Juntar TRCT, impugnar laudo)"],
          "base_legal": ["Artigo da CLT ou Constituição aplicado"],
          "jurisprudencia": ["Sempre use a tag de omissão ou copie a ementa real fornecida nos autos."],
          "doutrina": ["Tese doutrinária de apoio."]
        }
        """
        
        prompt_mestre = f"DIRECIONAMENTO:\n{fatos}\n\n--- DADOS COLETADOS ---\n{texto_consolidado}\n\nGere o Dossiê Trabalhista em formato JSON estrito. Se a informação não existir, responda 'INFORMAÇÃO NÃO ENCONTRADA' nos campos correspondentes."

        config_mestre = types.GenerateContentConfig(
            system_instruction=instrucao_sistema_mestre, 
            temperature=0.0, 
            response_mime_type="application/json",
            tools=[{"googleSearch": {}}],
            safety_settings=filtros_seguranca
        )

        response_final = client.models.generate_content(model='gemini-2.5-flash', contents=[prompt_mestre], config=config_mestre)

        if not hasattr(response_final, 'text') or not response_final.text: raise Exception("Recusado pelos filtros.")

        # ORDEM DE SEGURANÇA 2.0: JSON PRIMEIRO, ESCUDO DEPOIS
        dados_brutos = json.loads(response_final.text.strip(), strict=False)
        dados = aplicar_escudo_no_json(dados_brutos)
        callback_sucesso(dados)

    except Exception as e:
        callback_erro(str(e))

def gerar_dossie_word(dados, caminho_salvar):
    try:
        doc = docx.Document()
        for s in doc.sections: s.top_margin, s.bottom_margin, s.left_margin, s.right_margin = Cm(3), Cm(2), Cm(3), Cm(2)
        doc.add_heading('M.A | DOSSIÊ ESTRATÉGICO TRABALHISTA', level=1)
        
        def add_secao(titulo, conteudo):
            if conteudo:
                doc.add_heading(titulo, level=2)
                p = doc.add_paragraph(conteudo)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        add_secao('Resumo Estratégico e Probabilidade', dados.get('resumo_estrategico', ''))
        if dados.get('analise_foco_cirurgico'):
            doc.add_heading('Foco Cirúrgico (Auditoria Específica)', level=2)
            for item in dados.get('analise_foco_cirurgico', []): doc.add_paragraph(item, style='List Bullet')
        add_secao('Jurimetria', dados.get('jurimetria', ''))
        add_secao('Report para o Cliente', dados.get('resumo_cliente', ''))

        if dados.get('timeline'):
            doc.add_heading('Cronologia Contratual (Admissão/Demissão)', level=2)
            for t in dados.get('timeline', []): doc.add_paragraph(f"{t.get('data', '')}: {t.get('evento', '')}", style='List Bullet')
        if dados.get('vulnerabilidades_contraparte'):
            doc.add_heading('Modo Combate (Falhas Probatórias e Prescrição)', level=2)
            for v in dados.get('vulnerabilidades_contraparte', []): doc.add_paragraph(v, style='List Bullet')
        if dados.get('checklist'):
            doc.add_heading('Providências da Defesa', level=2)
            for c in dados.get('checklist', []): doc.add_paragraph(c, style='List Bullet')
        if dados.get('base_legal'):
            doc.add_heading('Base Legal (CLT e Leis Especiais)', level=2)
            for b in dados.get('base_legal', []): doc.add_paragraph(b, style='List Bullet')
        if dados.get('jurisprudencia'):
            doc.add_heading('Jurisprudência Localizada', level=2)
            for j in dados.get('jurisprudencia', []): doc.add_paragraph(j, style='List Bullet')
        if dados.get('doutrina'):
            doc.add_heading('Doutrina Trabalhista', level=2)
            for d in dados.get('doutrina', []): doc.add_paragraph(d, style='List Bullet')

        doc.save(caminho_salvar)
        return True
    except: return False

janela = tk.Tk()
janela.title("M.A | Auditor Trabalhista (Blindado)")
janela.geometry("1100x750")
janela.configure(bg="#020617")
dados_analise_atual = {}

style = ttk.Style()
style.theme_use('clam')
style.configure("TNotebook", background="#020617", borderwidth=0)
style.configure("TNotebook.Tab", background="#0f172a", foreground="#94a3b8", padding=[20, 10], font=("Segoe UI", 10, "bold"), borderwidth=0)
# Laranja Trabalhista
style.map("TNotebook.Tab", background=[("selected", "#c2410c")], foreground=[("selected", "white")])

def executar_analise():
    fatos = txt_fatos.get(1.0, tk.END).strip()
    caminhos_txt = filedialog.askopenfilenames(filetypes=[("Arquivos de Texto (Volumes)", "*.txt")])
    if not caminhos_txt: return
    caminhos_txt = sorted(list(caminhos_txt))

    btn_executar.config(state=tk.DISABLED, text="⏳ AUDITANDO VOLUMES COM FILTRO REGEX...", bg="#64748b")
    progress_bar.pack(side=tk.LEFT, padx=10)
    progress_bar['maximum'] = len(caminhos_txt)
    progress_bar['value'] = 0
    
    def atualizar_progresso(atual, total, mensagem):
        def _ui_update():
            lbl_status.config(text=mensagem)
            progress_bar['value'] = atual
        janela.after(0, _ui_update)

    def sucesso(dados):
        def _update_ui_sucesso():
            progress_bar.pack_forget() 
            global dados_analise_atual
            dados_analise_atual = dados
            
            txt_resumo.delete(1.0, tk.END)
            txt_resumo.insert(tk.END, f"=== ESTRATÉGIA ===\n{dados.get('resumo_estrategico','')}\n\n=== JURIMETRIA ===\n{dados.get('jurimetria','')}\n\n=== REPORT CLIENTE ===\n{dados.get('resumo_cliente','')}")
            txt_foco.delete(1.0, tk.END)
            txt_foco.insert(tk.END, "=== ANÁLISE CIRÚRGICA ===\n\n" + "\n\n".join([f"📌 {a}" for a in dados.get('analise_foco_cirurgico', [])]))
            txt_vuln.delete(1.0, tk.END)
            txt_vuln.insert(tk.END, "=== FALHAS E PRESCRIÇÃO ===\n" + "\n".join([f"- {v}" for v in dados.get('vulnerabilidades_contraparte', [])]) + "\n\n")
            txt_vuln.insert(tk.END, "=== TIMELINE ===\n" + "\n".join([f"[{t.get('data', '')}] {t.get('evento', '')}" for t in dados.get('timeline', [])]))
            txt_legal.delete(1.0, tk.END)
            txt_legal.insert(tk.END, "=== JURISPRUDÊNCIA (BLINDADA) ===\n" + "\n\n".join([f"- {j}" for j in dados.get('jurisprudencia', [])]) + "\n\n")
            txt_legal.insert(tk.END, "=== BASE LEGAL TRABALHISTA ===\n" + "\n".join([f"- {b}" for b in dados.get('base_legal', [])]) + "\n\n")

            btn_executar.config(state=tk.NORMAL, text="🚀 ANALISAR NOVO LOTE DE VOLUMES", bg="#c2410c")
            lbl_status.config(text="Status: Auditoria concluída e blindada com sucesso.", fg="#10b981")
            notebook.select(2)
        janela.after(0, _update_ui_sucesso)

    def erro(msg):
        def _update_ui_erro():
            progress_bar.pack_forget() 
            btn_executar.config(state=tk.NORMAL, text="🚀 TENTAR NOVAMENTE", bg="#c2410c")
            lbl_status.config(text="Status: Erro na análise.", fg="#ef4444")
            messagebox.showerror("Erro", f"Falha ao auditar:\n{msg}")
        janela.after(0, _update_ui_erro)

    threading.Thread(target=analisar_lote_ia, args=(caminhos_txt, fatos, atualizar_progresso, sucesso, erro), daemon=True).start()

def salvar_dossie():
    if not dados_analise_atual: return messagebox.showinfo("Aviso", "Faça uma análise primeiro.")
    
    texto_total = json.dumps(dados_analise_atual)
    if "[" in texto_total or "]" in texto_total:
        resposta = messagebox.askyesno("Atenção", "O Dossiê possui marcações de colchetes '[ ]' que indicam informações a preencher. Deseja salvar o Word mesmo assim?")
        if not resposta: return

    caminho = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Document", "*.docx")], initialfile=f"Dossie_Trabalhista_Blindado_MA.docx")
    if caminho:
        gerar_dossie_word(dados_analise_atual, caminho)
        messagebox.showinfo("Sucesso", "Dossiê Trabalhista salvo com sucesso!")

sidebar = tk.Frame(janela, bg="#0f172a", width=300); sidebar.pack(side=tk.LEFT, fill=tk.Y); sidebar.pack_propagate(False)
tk.Label(sidebar, text="M.A | TRABALHISTA", font=("Segoe UI", 18, "bold"), bg="#0f172a", fg="#c2410c", pady=30).pack()
tk.Label(sidebar, text="Especialista em rescisões,\nhoras extras e prescrição.\n[Proteção Regex Ativa]", bg="#0f172a", fg="#cbd5e1", justify="center").pack(pady=10)
tk.Button(sidebar, text="📥 BAIXAR DOSSIÊ (WORD)", command=salvar_dossie, bg="#10b981", fg="white", font=("Segoe UI", 10, "bold"), pady=10).pack(side=tk.BOTTOM, fill="x", padx=20, pady=20)

main_area = tk.Frame(janela, bg="#020617"); main_area.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True, padx=20, pady=20)
notebook = ttk.Notebook(main_area); notebook.pack(fill=tk.BOTH, expand=True, pady=(0, 15))

tab_fatos = tk.Frame(notebook, bg="#1e293b"); notebook.add(tab_fatos, text=" 1. DIRECIONAMENTO ")
txt_fatos = scrolledtext.ScrolledText(tab_fatos, bg="#1e293b", fg="white", insertbackground="white", borderwidth=0, font=("Segoe UI", 11), padx=15, pady=15)
txt_fatos.pack(fill=tk.BOTH, expand=True); txt_fatos.insert(tk.END, "Ex: Atuamos para a empresa (Reclamada). Foque em invalidar as horas extras e buscar prescrição quinquenal...")

tab_res = tk.Frame(notebook, bg="#1e293b"); notebook.add(tab_res, text=" 2. ESTRATÉGIA ")
txt_resumo = scrolledtext.ScrolledText(tab_res, bg="#1e293b", fg="#e2e8f0", borderwidth=0, font=("Segoe UI", 11), padx=15, pady=15); txt_resumo.pack(fill=tk.BOTH, expand=True)

tab_foco = tk.Frame(notebook, bg="#1e293b"); notebook.add(tab_foco, text=" 3. FOCO CIRÚRGICO ")
txt_foco = scrolledtext.ScrolledText(tab_foco, bg="#1e293b", fg="#fbbf24", borderwidth=0, font=("Segoe UI", 11), padx=15, pady=15); txt_foco.pack(fill=tk.BOTH, expand=True)

tab_vuln = tk.Frame(notebook, bg="#1e293b"); notebook.add(tab_vuln, text=" 4. FALHAS E PRESCRIÇÃO ")
txt_vuln = scrolledtext.ScrolledText(tab_vuln, bg="#1e293b", fg="#fb923c", borderwidth=0, font=("Segoe UI", 11), padx=15, pady=15); txt_vuln.pack(fill=tk.BOTH, expand=True)

tab_leg = tk.Frame(notebook, bg="#1e293b"); notebook.add(tab_leg, text=" 5. FUNDAMENTAÇÃO ")
txt_legal = scrolledtext.ScrolledText(tab_leg, bg="#1e293b", fg="#34d399", borderwidth=0, font=("Segoe UI", 11), padx=15, pady=15); txt_legal.pack(fill=tk.BOTH, expand=True)

action_frame = tk.Frame(main_area, bg="#020617"); action_frame.pack(fill=tk.X, side=tk.BOTTOM)
btn_executar = tk.Button(action_frame, text="🚀 ANEXAR VOLUMES E AUDITAR (COM TRAVA REGEX)", command=executar_analise, bg="#c2410c", fg="white", font=("Segoe UI", 11, "bold"), pady=12, cursor="hand2")
btn_executar.pack(side=tk.LEFT, fill=tk.X, expand=True)

progress_bar = ttk.Progressbar(action_frame, mode='determinate', length=200)
lbl_status = tk.Label(action_frame, text="Status: Aguardando...", bg="#020617", fg="#64748b", font=("Segoe UI", 10)); lbl_status.pack(side=tk.LEFT, padx=20)

janela.mainloop()
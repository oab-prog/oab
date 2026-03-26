import os
import json
import re
import time
import threading
import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext, ttk
import docx
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from google import genai
from google.genai import types

# Tenta carregar o buscador web, se não tiver instalado, não trava o programa
try:
    from duckduckgo_search import DDGS
    HAS_DDGS = True
except ImportError:
    HAS_DDGS = False

API_KEY_LOCAL = os.getenv("GEMINI_API_KEY") 
dados_analise_atual = ""

# ==========================================
# ESCUDO 2.0 — saneamento recursivo de valores em JSON
# ==========================================
_PADRAO_JURIS_ESCUDO = re.compile(
    r'(?i)\b(?:HC|REsp|RHC|AgRg|AREsp|Apelação|Agravo|AI|Processo n[º°]?\s*)\s*[\d\.\-\/]+\/[A-Z]{2}\b'
)

def aplicar_escudo_no_json(dados):
    if isinstance(dados, dict):
        return {k: aplicar_escudo_no_json(v) for k, v in dados.items()}
    if isinstance(dados, list):
        return [aplicar_escudo_no_json(item) for item in dados]
    if isinstance(dados, str):
        return _PADRAO_JURIS_ESCUDO.sub(
            "[AVISO: NUMERAÇÃO OMITIDA - BUSQUE NA INTEGRA NO JUSBRASIL]", dados
        )
    return dados


def _laudo_json_para_markdown(d):
    """Converte o laudo JSON (pós-escudo) em texto com headings para UI e exportação Word."""
    sec = [
        ("RESUMO EXECUTIVO", "resumo_executivo"),
        ("RISCOS E CLÁUSULAS ABUSIVAS", "riscos_e_clausulas_abusivas"),
        ("OMISSÕES GRAVES", "omissoes_graves"),
        ("SUGESTÕES DE REDAÇÃO", "sugestoes_de_redacao"),
        ("FUNDAMENTAÇÃO LEGAL E JURISPRUDENCIAL", "fundamentacao_legal_e_jurisprudencial"),
    ]
    partes = []
    for titulo, chave in sec:
        val = d.get(chave)
        if val is None or (isinstance(val, str) and not val.strip()):
            val = "INFORMAÇÃO NÃO ENCONTRADA"
        partes.append(f"## {titulo}\n\n{val}")
    return "\n\n".join(partes)

def buscar_jurisprudencia_real(foco_busca):
    """Pesquisa silenciosamente na internet por julgados reais antes de acionar a IA."""
    if not HAS_DDGS: return "Aviso: Módulo de busca online ausente. (pip install duckduckgo-search)"
    try:
        termo = f"{foco_busca} nulidade cláusula abusiva contrato jurisprudencia site:jusbrasil.com.br"
        resultados = DDGS().text(termo, max_results=3)
        texto_pesquisa = "\n".join([f"- {r['title']}: {r['body']}" for r in resultados])
        return texto_pesquisa if texto_pesquisa else "Nenhuma jurisprudência específica encontrada online."
    except Exception:
        return "Busca online indisponível no momento."

# ==========================================
# LÓGICA DO REVISOR DE CONTRATOS
# ==========================================
def analisar_contrato(caminho_pdf, foco_usuario, callback_progresso, callback_sucesso, callback_erro):
    try:
        client = genai.Client(api_key=API_KEY_LOCAL.strip())
        filtros_seguranca = [
            types.SafetySetting(category="HARM_CATEGORY_HATE_SPEECH", threshold="BLOCK_NONE"),
            types.SafetySetting(category="HARM_CATEGORY_HARASSMENT", threshold="BLOCK_NONE"),
            types.SafetySetting(category="HARM_CATEGORY_SEXUALLY_EXPLICIT", threshold="BLOCK_NONE"),
            types.SafetySetting(category="HARM_CATEGORY_DANGEROUS_CONTENT", threshold="BLOCK_NONE"),
        ]

        # 1. Pesquisa Web Prévida (Contexto do Mundo Real)
        callback_progresso(10, "Buscando jurisprudência real na internet (Jusbrasil/STJ)...")
        pesquisa_web = buscar_jurisprudencia_real(foco_usuario)

        # 2. Upload Nativo do PDF (Lê inclusive contratos escaneados e assinados a caneta)
        callback_progresso(30, "Fazendo upload seguro e leitura óptica do contrato...")
        gemini_file = client.files.upload(file=caminho_pdf, config={'mime_type': 'application/pdf'})
        
        while True:
            f_info = client.files.get(name=gemini_file.name)
            if "FAILED" in str(f_info.state).upper(): raise Exception("Falha ao processar o PDF do contrato.")
            if "ACTIVE" in str(f_info.state).upper(): break
            time.sleep(2)

        # 3. Análise da Inteligência
        callback_progresso(60, "Auditando cláusulas, multas e brechas contratuais...")
        
        instrucao_sistema = """
        Você é o M.A | REVISOR DE CONTRATOS DE ELITE. Especialista em Direito Contratual e Empresarial.
        Sua missão é escanear o contrato fornecido, encontrar armadilhas, cláusulas abusivas, nulidades e sugerir melhorias para proteger o cliente.

        >>> VETO (Zero Hallucination 2.0) <<<
        Se o dado não existir, responda: INFORMAÇÃO NÃO ENCONTRADA

        >>> TRAVA ABSOLUTA ANTI-ALUCINAÇÃO (TOLERÂNCIA ZERO) <<<
        1. É TERMINANTEMENTE PROIBIDO inventar cláusulas, valores, partes ou prazos que não existam no documento PDF.
        2. Não invente números de processos, Recursos Especiais ou Súmulas.

        >>> SAÍDA OBRIGATÓRIA (JSON) <<<
        Responda SOMENTE um objeto JSON válido, sem markdown, com exatamente estas chaves (strings):
        - "resumo_executivo": quem são as partes, objeto, valor e prazo (com base no PDF).
        - "riscos_e_clausulas_abusivas": destaque o número da cláusula exata e explique o perigo.
        - "omissoes_graves": o que deveria estar no contrato para proteger o cliente, mas não está.
        - "sugestoes_de_redacao": como reescrever cláusulas problemáticas.
        - "fundamentacao_legal_e_jurisprudencial": use as informações da Pesquisa Web fornecida no prompt; não invente julgados.

        Em cada campo, use texto corrido ou parágrafos separados por \\n. Não inclua blocos de código nem ```json.
        """
        
        prompt = f"FOCO DO CLIENTE: {foco_usuario}\n\n--- JURISPRUDÊNCIA COLETADA DA WEB ---\n{pesquisa_web}\n\nAnalise o contrato em anexo e emita o Laudo de Revisão."
        
        config_revisor = types.GenerateContentConfig(
            system_instruction=instrucao_sistema,
            temperature=0.0,
            response_mime_type="application/json",
            safety_settings=filtros_seguranca,
        )
        
        conteudo_pdf = types.Part.from_uri(file_uri=f_info.uri, mime_type='application/pdf')
        
        resposta = client.models.generate_content(model='gemini-2.5-flash', contents=[conteudo_pdf, prompt], config=config_revisor)
        
        # 4. Apaga o arquivo por sigilo contratual
        callback_progresso(90, "Apagando contrato dos servidores (Garantia de Sigilo)...")
        client.files.delete(name=gemini_file.name)

        if not hasattr(resposta, 'text') or not resposta.text:
            raise Exception("A IA recusou-se a gerar a revisão.")

        dados_laudo = json.loads(resposta.text)
        dados_laudo = aplicar_escudo_no_json(dados_laudo)
        texto_blindado = _laudo_json_para_markdown(dados_laudo)

        callback_sucesso(texto_blindado)

    except Exception as e:
        callback_erro(str(e))
        try:
            if 'gemini_file' in locals() and gemini_file: client.files.delete(name=gemini_file.name)
        except: pass

def gerar_laudo_word(texto_laudo, caminho_salvar):
    try:
        doc = docx.Document()
        for s in doc.sections: s.top_margin, s.bottom_margin, s.left_margin, s.right_margin = Cm(3), Cm(2), Cm(3), Cm(2)
        doc.add_heading('M.A | LAUDO TÉCNICO DE REVISÃO CONTRATUAL', level=1)
        
        linhas = texto_laudo.split('\n')
        for l in linhas:
            texto_limpo = l.replace("**", "").replace("*", "").strip()
            if not texto_limpo: continue
            
            if l.startswith("#"):
                nivel = l.count("#")
                doc.add_heading(texto_limpo.replace("#", "").strip(), level=min(nivel, 3))
            elif l.startswith("- ") or l.startswith("* "):
                doc.add_paragraph(texto_limpo[1:].strip(), style='List Bullet')
            else:
                p = doc.add_paragraph(texto_limpo)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        doc.save(caminho_salvar)
        return True
    except: return False

# ==========================================
# INTERFACE GRÁFICA DO REVISOR
# ==========================================
janela = tk.Tk()
janela.title("M.A | Revisor de Contratos (Visão Computacional e Web)")
janela.geometry("1000x700")
janela.configure(bg="#020617")

style = ttk.Style()
style.theme_use('clam')

def iniciar_revisao():
    foco = txt_foco.get(1.0, tk.END).strip()
    if not foco: return messagebox.showwarning("Aviso", "Preencha o foco da análise.")
    
    caminho_pdf = filedialog.askopenfilename(filetypes=[("PDF Files", "*.pdf")])
    if not caminho_pdf: return

    btn_executar.config(state=tk.DISABLED, text="⏳ ANALISANDO CONTRATO...", bg="#64748b")
    progress_bar.pack(fill=tk.X, pady=(10, 0))
    progress_bar['value'] = 0
    
    def atualizar_progresso(valor, mensagem):
        def _ui_update():
            lbl_status.config(text=mensagem)
            progress_bar['value'] = valor
        janela.after(0, _ui_update)

    def sucesso(laudo):
        def _update_ui_sucesso():
            progress_bar.pack_forget() 
            global dados_analise_atual
            dados_analise_atual = laudo
            
            txt_resultado.delete(1.0, tk.END)
            txt_resultado.insert(tk.END, laudo)

            btn_executar.config(state=tk.NORMAL, text="🔍 SELECIONAR NOVO CONTRATO E REVISAR", bg="#2563eb")
            lbl_status.config(text="Status: Revisão contratual concluída com sucesso.", fg="#10b981")
        janela.after(0, _update_ui_sucesso)

    def erro(msg):
        def _update_ui_erro():
            progress_bar.pack_forget() 
            btn_executar.config(state=tk.NORMAL, text="🔍 TENTAR NOVAMENTE", bg="#2563eb")
            lbl_status.config(text="Status: Erro na revisão.", fg="#ef4444")
            messagebox.showerror("Erro", f"Falha ao revisar o contrato:\n{msg}")
        janela.after(0, _update_ui_erro)

    threading.Thread(target=analisar_contrato, args=(caminho_pdf, foco, atualizar_progresso, sucesso, erro), daemon=True).start()

def salvar_laudo():
    if not dados_analise_atual: return messagebox.showinfo("Aviso", "Faça uma revisão primeiro.")
    caminho = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Document", "*.docx")], initialfile=f"Revisao_Contratual_MA.docx")
    if caminho:
        gerar_laudo_word(dados_analise_atual, caminho)
        messagebox.showinfo("Sucesso", "Laudo Contratual salvo com sucesso!")

# --- Layout Visual ---
sidebar = tk.Frame(janela, bg="#0f172a", width=280)
sidebar.pack(side=tk.LEFT, fill=tk.Y); sidebar.pack_propagate(False)

tk.Label(sidebar, text="M.A | REVISOR", font=("Segoe UI", 18, "bold"), bg="#0f172a", fg="#3b82f6", pady=30).pack()
tk.Label(sidebar, text="Caçador de armadilhas,\nnulidades e cláusulas\nabusivas em contratos.", bg="#0f172a", fg="#cbd5e1", justify="center").pack(pady=10)
tk.Label(sidebar, text="[ OCR Nativo Ativado ]\n[ Busca Web DuckDuckGo ]\n[ Escudo 2.0 JSON ]", bg="#0f172a", fg="#64748b", justify="center", font=("Segoe UI", 9, "italic")).pack(pady=20)
tk.Button(sidebar, text="📥 BAIXAR LAUDO (WORD)", command=salvar_laudo, bg="#10b981", fg="white", font=("Segoe UI", 10, "bold"), pady=10).pack(side=tk.BOTTOM, fill="x", padx=20, pady=20)

main_area = tk.Frame(janela, bg="#020617", padx=20, pady=20); main_area.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True)

tk.Label(main_area, text=" PASSO 1: Foco da Análise e Diretriz para o M.A", font=("Segoe UI", 10, "bold"), bg="#020617", fg="#94a3b8", anchor="w").pack(fill=tk.X, pady=5)
txt_foco = scrolledtext.ScrolledText(main_area, height=3, bg="#1e293b", fg="white", borderwidth=0, font=("Segoe UI", 11), padx=10, pady=10)
txt_foco.pack(fill=tk.X, pady=5)
txt_foco.insert(tk.END, "Proteger os direitos do Cliente contra cláusulas abusivas, juros extorsivos e riscos de quebra de contrato. Encontrar brechas de rescisão.")

action_frame = tk.Frame(main_area, bg="#020617"); action_frame.pack(fill=tk.X, pady=(15, 10))
btn_executar = tk.Button(action_frame, text="🔍 SELECIONAR CONTRATO EM PDF E REVISAR", command=iniciar_revisao, bg="#2563eb", fg="white", font=("Segoe UI", 12, "bold"), pady=12, cursor="hand2")
btn_executar.pack(side=tk.LEFT, fill=tk.X, expand=True)

progress_bar = ttk.Progressbar(main_area, mode='determinate', length=200)
lbl_status = tk.Label(main_area, text="Status: Aguardando...", bg="#020617", fg="#64748b", font=("Segoe UI", 10)); lbl_status.pack(anchor="w", pady=5)

tk.Label(main_area, text=" PASSO 2: Laudo Pericial Contratual", font=("Segoe UI", 10, "bold"), bg="#020617", fg="#94a3b8", anchor="w").pack(fill=tk.X, pady=(15, 5))
txt_resultado = scrolledtext.ScrolledText(main_area, bg="#f8fafc", fg="#0f172a", borderwidth=0, font=("Segoe UI", 11), padx=15, pady=15)
txt_resultado.pack(fill=tk.BOTH, expand=True)

janela.mainloop()
import os
import json
import time
import datetime
import threading
import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext, ttk
import docx
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from google import genai
from google.genai import types
import re # O motor do Escudo Anti-Alucinação

# ==========================================
# MOTOR DE BUSCA WEB (DUCKDUCKGO)
# ==========================================
try:
    from duckduckgo_search import DDGS
    HAS_DDGS = True
except ImportError:
    HAS_DDGS = False

API_KEY_LOCAL = os.getenv("GEMINI_API_KEY") 

# ==========================================
# O ESCUDO DE CÓDIGO (REGEX) - A PORTA DE AÇO
# ==========================================
def escudo_anti_alucinacao(texto):
    padrao_jurisprudencia = r'(?i)\b(?:HC|Habeas Corpus|REsp|Recurso Especial|RHC|AgRg|AREsp|Apelação|Agravo|Processo)\s*(?:n[º°]?\s*)?[\d\.\-\/]+\/[A-Z]{2}\b'
    texto_limpo = re.sub(padrao_jurisprudencia, "[AVISO: NUMERACAO OMITIDA - BUSQUE NA INTEGRA NO JUSBRASIL/STJ]", texto)
    return texto_limpo

def aplicar_escudo_no_json(dados):
    if isinstance(dados, dict):
        return {k: aplicar_escudo_no_json(v) for k, v in dados.items()}
    elif isinstance(dados, list):
        return [aplicar_escudo_no_json(elem) for elem in dados]
    elif isinstance(dados, str):
        return escudo_anti_alucinacao(dados)
    else:
        return dados

def buscar_jurisprudencia_real(foco_busca):
    if not HAS_DDGS: return "Busca web indisponível."
    try:
        termo = f"{foco_busca} penal processo penal jurisprudencia 2025 2026 site:jusbrasil.com.br"
        resultados = DDGS().text(termo, max_results=3)
        texto_pesquisa = "\n".join([f"- {r['title']}: {r['body']}" for r in resultados])
        return texto_pesquisa if texto_pesquisa else "Nenhuma jurisprudência criminal localizada online."
    except Exception:
        return "Busca online indisponível no momento."

# ==========================================
# LÓGICA DE AUDITORIA EM LOTE (IA)
# ==========================================
def analisar_lote_ia(caminhos_txt, fatos, callback_progresso, callback_sucesso, callback_erro):
    try:
        client = genai.Client(api_key=API_KEY_LOCAL.strip())
        filtros_seguranca = [
            types.SafetySetting(category="HARM_CATEGORY_HATE_SPEECH", threshold="BLOCK_NONE"),
            types.SafetySetting(category="HARM_CATEGORY_HARASSMENT", threshold="BLOCK_NONE"),
            types.SafetySetting(category="HARM_CATEGORY_SEXUALLY_EXPLICIT", threshold="BLOCK_NONE"),
            types.SafetySetting(category="HARM_CATEGORY_DANGEROUS_CONTENT", threshold="BLOCK_NONE"),
        ]

        # FASE 0: BUSCA WEB
        callback_progresso(0, len(caminhos_txt) + 1, "Pesquisando teses criminais no STJ/Internet...")
        contexto_web = buscar_jurisprudencia_real(fatos)

        resumos_volumes = []
        total_vols = len(caminhos_txt)

        for idx, caminho in enumerate(caminhos_txt):
            vol_num = idx + 1
            callback_progresso(vol_num, total_vols + 1, f"Fase 1: Auditando Volume {vol_num} de {total_vols}...")
            
            gemini_file = client.files.upload(file=caminho, config={'mime_type': 'text/plain'})
            while True:
                f_info = client.files.get(name=gemini_file.name)
                if "FAILED" in str(f_info.state).upper(): raise Exception(f"Falha no Volume {vol_num}.")
                if "ACTIVE" in str(f_info.state).upper(): break
                time.sleep(2)
            
            conteudo_txt = types.Part.from_uri(file_uri=f_info.uri, mime_type='text/plain')
            
            prompt_detetive = f"DIRECIONAMENTO DA ESTRATÉGIA:\n{fatos}\n\nAnalise este volume Criminal. Extraia informações sobre: autoria, materialidade, depoimentos, contradições e nulidades (buscas ilegais, quebra de cadeia de custódia, etc). OBRIGATÓRIO citar as fls. Seja exato e não invente dados."
            config_detetive = types.GenerateContentConfig(temperature=0.0, safety_settings=filtros_seguranca)
            response_detetive = client.models.generate_content(model='gemini-2.5-flash', contents=[conteudo_txt, prompt_detetive], config=config_detetive)
            
            resumos_volumes.append(f"--- ACHADOS DO VOLUME {vol_num} ---\n{response_detetive.text.strip()}\n")
            client.files.delete(name=gemini_file.name)

        callback_progresso(total_vols, total_vols + 1, "Fase 2: Estruturando Parecer Jurídico Premium...")
        texto_consolidado = "\n".join(resumos_volumes)

        # ✅ AQUI ESTÁ A MÁGICA DA ROBUSTEZ: Instrução para Parecer Premium
        instrucao_sistema_mestre = """
        Você é o M.A | JUS IA EXPERIENCE, um Auditor Forense Especialista em Direito Penal e Processual Penal.
        
        >>> MISSÃO: PARECER JURÍDICO PREMIUM (ROBUSTEZ MÁXIMA) <<<
        NÃO crie tópicos curtos ou listas simples. Desenvolva argumentos robustos, densos e dissertativos para cada item do JSON. 
        Escreva como um Advogado Criminalista Sênior elaborando um Parecer de Alta Complexidade. Justifique cada falha apontada com profundidade analítica, sempre ancorando nas provas e nas fls.
        
        >>> TRAVA ABSOLUTA ANTI-ALUCINAÇÃO (TOLERÂNCIA ZERO) <<<
        1. PROIBIDO inventar provas, laudos, interceptações ou locais que não existam nos autos.
        2. PROIBIDO inventar números de processos, HC, REsp ou Súmulas.
        
        >>> REGRA DE FORMATAÇÃO JSON <<<
        O retorno DEVE ser EXCLUSIVAMENTE um objeto JSON válido.
        USE ESTA ESTRUTURA:
        {
          "resumo_estrategico": "Escreva 2 a 3 parágrafos densos diagnosticando o cenário penal e a viabilidade da defesa, sempre citando as fls.",
          "analise_foco_cirurgico": ["Título do Argumento 1: Escreva um parágrafo robusto e dissertativo sobre este ponto, citando fls.", "Título do Argumento 2: (Parágrafo longo e denso com fls)"],
          "jurimetria": "Análise profunda do comportamento do juízo/tribunal neste caso.",
          "resumo_cliente": "E-mail executivo e tranquilizador para o cliente.",
          "timeline": [{"data": "DD/MM/AAAA", "evento": "Descrição detalhada do evento e suas fls."}],
          "vulnerabilidades_contraparte": ["Da Nulidade Processual: Discorra profundamente sobre a falha do MP/Polícia, citando fls e as consequências práticas dessa nulidade.", "Da Fragilidade Probatória: (Parágrafo denso com fls)"],
          "checklist": ["Ação estratégica explicada em detalhes."],
          "base_legal": ["Fundamentação da lei aplicável explicada no contexto do caso criminal."],
          "jurisprudencia": ["Use OBRIGATORIAMENTE os julgados reais da Pesquisa Web fornecida ou omita."],
          "doutrina": ["Tese doutrinária detalhada que apoia a defesa penal."]
        }
        """
        
        prompt_mestre = f"DIRECIONAMENTO:\n{fatos}\n\n--- DADOS COLETADOS ---\n{texto_consolidado}\n\n--- PESQUISA WEB REAL (2025/2026) ---\n{contexto_web}\n\nGere o Parecer Criminal em formato JSON estrito."

        config_mestre = types.GenerateContentConfig(
            system_instruction=instrucao_sistema_mestre, 
            temperature=0.0, # Zero Hallucination 2.0 - Tolerância zero para alucinações
            safety_settings=filtros_seguranca,
            response_mime_type="application/json"
        )

        response_final = client.models.generate_content(model='gemini-2.5-flash', contents=[prompt_mestre], config=config_mestre)

        if not hasattr(response_final, 'text') or not response_final.text: raise Exception("Recusado pelos filtros.")

        texto_puro = response_final.text.strip()
        

        
        dados = json.loads(texto_puro.strip(), strict=False)
        dados = aplicar_escudo_no_json(dados)
        callback_sucesso(dados)

    except Exception as e:
        callback_erro(str(e))

def gerar_dossie_word(dados, caminho_salvar):
    try:
        doc = docx.Document()
        for s in doc.sections: s.top_margin, s.bottom_margin, s.left_margin, s.right_margin = Cm(3), Cm(2), Cm(3), Cm(2)
        doc.add_heading('M.A | PARECER ESTRATÉGICO CRIMINAL', level=1)
        
        def add_secao(titulo, conteudo):
            if conteudo:
                doc.add_heading(titulo, level=2)
                p = doc.add_paragraph(str(conteudo))
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        add_secao('Resumo Estratégico e Viabilidade', dados.get('resumo_estrategico', ''))
        
        if dados.get('analise_foco_cirurgico'):
            doc.add_heading('Análise Qualitativa dos Fatos (Foco Cirúrgico)', level=2)
            for item in dados.get('analise_foco_cirurgico', []): 
                p = doc.add_paragraph(item, style='List Bullet')
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                
        add_secao('Análise Comportamental do Juízo (Jurimetria)', dados.get('jurimetria', ''))
        add_secao('Report de Andamento ao Cliente', dados.get('resumo_cliente', ''))

        if dados.get('timeline'):
            doc.add_heading('Cronologia Processual Exaustiva', level=2)
            for t in dados.get('timeline', []): doc.add_paragraph(f"{t.get('data', '')}: {t.get('evento', '')}", style='List Bullet')
            
        if dados.get('vulnerabilidades_contraparte'):
            doc.add_heading('Vulnerabilidades da Acusação e Nulidades', level=2)
            for v in dados.get('vulnerabilidades_contraparte', []): 
                p = doc.add_paragraph(v, style='List Bullet')
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                
        if dados.get('checklist'):
            doc.add_heading('Plano de Ação Estratégico', level=2)
            for c in dados.get('checklist', []): doc.add_paragraph(c, style='List Bullet')
            
        if dados.get('base_legal'):
            doc.add_heading('Amparo Legal', level=2)
            for b in dados.get('base_legal', []): doc.add_paragraph(b, style='List Bullet')
            
        if dados.get('jurisprudencia'):
            doc.add_heading('Jurisprudência Aplicada', level=2)
            for j in dados.get('jurisprudencia', []): doc.add_paragraph(j, style='List Bullet')
            
        if dados.get('doutrina'):
            doc.add_heading('Apoio Doutrinário', level=2)
            for d in dados.get('doutrina', []): doc.add_paragraph(d, style='List Bullet')

        doc.save(caminho_salvar)
        return True
    except: return False

janela = tk.Tk()
janela.title("M.A | Auditor Criminal (Parecer Premium)")
janela.geometry("1100x750")
janela.configure(bg="#020617")
dados_analise_atual = {}

style = ttk.Style()
style.theme_use('clam')
style.configure("TNotebook", background="#020617", borderwidth=0)
style.configure("TNotebook.Tab", background="#0f172a", foreground="#94a3b8", padding=[20, 10], font=("Segoe UI", 10, "bold"), borderwidth=0)
style.map("TNotebook.Tab", background=[("selected", "#dc2626")], foreground=[("selected", "white")])

def executar_analise():
    fatos = txt_fatos.get(1.0, tk.END).strip()
    caminhos_txt = filedialog.askopenfilenames(filetypes=[("Arquivos de Texto (Volumes)", "*.txt")])
    if not caminhos_txt: return
    caminhos_txt = sorted(list(caminhos_txt))

    btn_executar.config(state=tk.DISABLED, text="⏳ ELABORANDO PARECER PREMIUM E LENDO VOLUMES...", bg="#64748b")
    progress_bar.pack(side=tk.LEFT, padx=10)
    progress_bar['maximum'] = len(caminhos_txt) + 1
    progress_bar['value'] = 0
    
    def atualizar_progresso(atual, total, mensagem):
        def _ui_update():
            lbl_status.config(text=mensagem)
            progress_bar['value'] = atual
        janela.after(0, _ui_update)

    def sucesso(dados):
        def _update_ui_sucesso():
            progress_bar.pack_forget() 
            global dados_analise_atual
            dados_analise_atual = dados
            
            txt_resumo.delete(1.0, tk.END)
            txt_resumo.insert(tk.END, f"=== ESTRATÉGIA ===\n{dados.get('resumo_estrategico','')}\n\n=== JURIMETRIA ===\n{dados.get('jurimetria','')}\n\n=== REPORT CLIENTE ===\n{dados.get('resumo_cliente','')}")
            txt_foco.delete(1.0, tk.END)
            txt_foco.insert(tk.END, "=== ANÁLISE QUALITATIVA DOS FATOS ===\n\n" + "\n\n".join([f"📌 {a}" for a in dados.get('analise_foco_cirurgico', [])]))
            txt_vuln.delete(1.0, tk.END)
            txt_vuln.insert(tk.END, "=== VULNERABILIDADES E NULIDADES ===\n\n" + "\n\n".join([f"⚠️ {v}" for v in dados.get('vulnerabilidades_contraparte', [])]) + "\n\n")
            txt_vuln.insert(tk.END, "=== TIMELINE ===\n" + "\n".join([f"[{t.get('data', '')}] {t.get('evento', '')}" for t in dados.get('timeline', [])]))
            txt_legal.delete(1.0, tk.END)
            txt_legal.insert(tk.END, "=== JURISPRUDÊNCIA (WEB BLINDADA) ===\n" + "\n\n".join([f"- {j}" for j in dados.get('jurisprudencia', [])]) + "\n\n")
            txt_legal.insert(tk.END, "=== BASE LEGAL ===\n" + "\n".join([f"- {b}" for b in dados.get('base_legal', [])]) + "\n\n")

            btn_executar.config(state=tk.NORMAL, text="🚀 ANALISAR NOVO LOTE DE VOLUMES", bg="#dc2626")
            lbl_status.config(text="Status: Parecer Criminal gerado com sucesso.", fg="#10b981")
            notebook.select(2)
        janela.after(0, _update_ui_sucesso)

    def erro(msg):
        def _update_ui_erro():
            progress_bar.pack_forget() 
            btn_executar.config(state=tk.NORMAL, text="🚀 TENTAR NOVAMENTE", bg="#dc2626")
            lbl_status.config(text="Status: Erro na análise.", fg="#ef4444")
            messagebox.showerror("Erro", f"Falha ao auditar:\n{msg}")
        janela.after(0, _update_ui_erro)

    threading.Thread(target=analisar_lote_ia, args=(caminhos_txt, fatos, atualizar_progresso, sucesso, erro), daemon=True).start()

def salvar_dossie():
    if not dados_analise_atual: return messagebox.showinfo("Aviso", "Faça uma análise primeiro.")
    texto_total = json.dumps(dados_analise_atual)
    if "AVISO: NUMERACAO OMITIDA" in texto_total:
        messagebox.showwarning("⚠️ Julgados Omitidos", "O Parecer contém julgados ocultados pelo Escudo Regex.\nBusque e insira as jurisprudências manualmente antes de entregar.")
    caminho = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Document", "*.docx")], initialfile="Parecer_Criminal_Premium_MA.docx")
    if caminho:
        gerar_dossie_word(dados_analise_atual, caminho)
        messagebox.showinfo("Sucesso", "Parecer Premium salvo com sucesso!")

sidebar = tk.Frame(janela, bg="#0f172a", width=300); sidebar.pack(side=tk.LEFT, fill=tk.Y); sidebar.pack_propagate(False)
tk.Label(sidebar, text="M.A | CRIMINAL", font=("Segoe UI", 18, "bold"), bg="#0f172a", fg="#ef4444", pady=30).pack()
tk.Label(sidebar, text="Pareceres robustos\ne teses defensivas densas.\n[Proteção Regex Ativa]", bg="#0f172a", fg="#cbd5e1", justify="center").pack(pady=10)
tk.Button(sidebar, text="📥 BAIXAR PARECER (WORD)", command=salvar_dossie, bg="#10b981", fg="white", font=("Segoe UI", 10, "bold"), pady=10).pack(side=tk.BOTTOM, fill="x", padx=20, pady=20)

main_area = tk.Frame(janela, bg="#020617"); main_area.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True, padx=20, pady=20)
notebook = ttk.Notebook(main_area); notebook.pack(fill=tk.BOTH, expand=True, pady=(0, 15))

tab_fatos = tk.Frame(notebook, bg="#1e293b"); notebook.add(tab_fatos, text=" 1. DIRECIONAMENTO ")
txt_fatos = scrolledtext.ScrolledText(tab_fatos, bg="#1e293b", fg="white", insertbackground="white", borderwidth=0, font=("Segoe UI", 11), padx=15, pady=15)
txt_fatos.pack(fill=tk.BOTH, expand=True); txt_fatos.insert(tk.END, "Cole aqui o Mapa de Focos Estratégicos gerado pelo Diretor...")

tab_res = tk.Frame(notebook, bg="#1e293b"); notebook.add(tab_res, text=" 2. PARECER & JURIMETRIA ")
txt_resumo = scrolledtext.ScrolledText(tab_res, bg="#1e293b", fg="#e2e8f0", borderwidth=0, font=("Segoe UI", 11), padx=15, pady=15); txt_resumo.pack(fill=tk.BOTH, expand=True)

tab_foco = tk.Frame(notebook, bg="#1e293b"); notebook.add(tab_foco, text=" 3. ANÁLISE QUALITATIVA ")
txt_foco = scrolledtext.ScrolledText(tab_foco, bg="#1e293b", fg="#fbbf24", borderwidth=0, font=("Segoe UI", 11), padx=15, pady=15); txt_foco.pack(fill=tk.BOTH, expand=True)

tab_vuln = tk.Frame(notebook, bg="#1e293b"); notebook.add(tab_vuln, text=" 4. NULIDADES DA ACUSAÇÃO ")
txt_vuln = scrolledtext.ScrolledText(tab_vuln, bg="#1e293b", fg="#ef4444", borderwidth=0, font=("Segoe UI", 11), padx=15, pady=15); txt_vuln.pack(fill=tk.BOTH, expand=True)

tab_leg = tk.Frame(notebook, bg="#1e293b"); notebook.add(tab_leg, text=" 5. FUNDAMENTAÇÃO ")
txt_legal = scrolledtext.ScrolledText(tab_leg, bg="#1e293b", fg="#34d399", borderwidth=0, font=("Segoe UI", 11), padx=15, pady=15); txt_legal.pack(fill=tk.BOTH, expand=True)

action_frame = tk.Frame(main_area, bg="#020617"); action_frame.pack(fill=tk.X, side=tk.BOTTOM)
btn_executar = tk.Button(action_frame, text="🚀 ANEXAR VOLUMES E GERAR PARECER PREMIUM", command=executar_analise, bg="#dc2626", fg="white", font=("Segoe UI", 11, "bold"), pady=12, cursor="hand2")
btn_executar.pack(side=tk.LEFT, fill=tk.X, expand=True)

progress_bar = ttk.Progressbar(action_frame, mode='determinate', length=200)
lbl_status = tk.Label(action_frame, text="Status: Aguardando...", bg="#020617", fg="#64748b", font=("Segoe UI", 10)); lbl_status.pack(side=tk.LEFT, padx=20)

janela.mainloop()
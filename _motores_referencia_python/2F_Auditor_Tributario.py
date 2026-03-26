import os
import json
import time
import datetime
import threading
import tkinter as tk
import re
from tkinter import filedialog, messagebox, scrolledtext, ttk
import docx
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from google import genai
from google.genai import types
from typing import Dict, List, Any, Optional

API_KEY_LOCAL = os.getenv("GEMINI_API_KEY") 

# ==========================================
# O ESCUDO DE CÓDIGO (REGEX) - TRIBUTÁRIO
# ==========================================
def aplicar_escudo_no_json(dados: Any) -> Any:
    """Sanitiza strings via Regex de forma recursiva (Zero Hallucination 2.0)."""
    padrao_jurisprudencia = r'(?i)\b(?:REsp|RE|RMS|AgRg|AREsp|Apelação|Agravo|Processo n[º°]?\s*)\s*[\d\.\-\/]+\/[A-Z]{2}\b'
    
    if isinstance(dados, dict):
        return {k: aplicar_escudo_no_json(v) for k, v in dados.items()}
    elif isinstance(dados, list):
        return [aplicar_escudo_no_json(i) for i in dados]
    elif isinstance(dados, str):
        return re.sub(padrao_jurisprudencia, "[AVISO: NUMERACAO OMITIDA - BUSQUE NO JUSBRASIL/STJ/STF]", dados)
    else:
        return dados

def analisar_lote_ia(caminhos_txt: List[str], fatos: str, callback_progresso: Any, callback_sucesso: Any, callback_erro: Any):
    try:
        client = genai.Client(api_key=API_KEY_LOCAL.strip())
        filtros_seguranca = [
            types.SafetySetting(category="HARM_CATEGORY_HATE_SPEECH", threshold="BLOCK_NONE"),
            types.SafetySetting(category="HARM_CATEGORY_HARASSMENT", threshold="BLOCK_NONE"),
            types.SafetySetting(category="HARM_CATEGORY_SEXUALLY_EXPLICIT", threshold="BLOCK_NONE"),
            types.SafetySetting(category="HARM_CATEGORY_DANGEROUS_CONTENT", threshold="BLOCK_NONE"),
        ]

        resumos_volumes = []
        total_vols = len(caminhos_txt)

        for idx, caminho in enumerate(caminhos_txt):
            vol_num = idx + 1
            callback_progresso(vol_num, total_vols, f"Fase 1: Auditando Créditos Tributários no Volume {vol_num} de {total_vols}...")
            
            gemini_file = client.files.upload(file=caminho, config={'mime_type': 'text/plain'})
            while True:
                f_info = client.files.get(name=gemini_file.name)
                if "FAILED" in str(f_info.state).upper(): raise Exception(f"Falha no Volume {vol_num}.")
                if "ACTIVE" in str(f_info.state).upper(): break
                time.sleep(2)
            
            conteudo_txt = types.Part.from_uri(file_uri=f_info.uri, mime_type='text/plain')
            
            prompt_detetive = f"DIRECIONAMENTO DA ESTRATÉGIA:\n{fatos}\n\nLeia este volume Tributário. Extraia lançamentos, notificações, CDAs, datas de prescrição/decadência e fatos geradores. É OBRIGATÓRIO citar as fls. de cada achado. RIGOR MÁXIMO para Mandado de Segurança (8x), Agravo (7x), Apelação (7x) e Repetição de Indébito (6x). NÃO INVENTE DADOS."
            config_detetive = types.GenerateContentConfig(temperature=0.0, safety_settings=filtros_seguranca)
            response_detetive = client.models.generate_content(model='gemini-2.0-flash', contents=[conteudo_txt, prompt_detetive], config=config_detetive)
            
            resumos_volumes.append(f"--- ACHADOS TRIBUTÁRIOS DO VOLUME {vol_num} ---\n{response_detetive.text.strip()}\n")
            client.files.delete(name=gemini_file.name)

        callback_progresso(total_vols, total_vols, "Fase 2: Estruturando Dossiê Jurídico Tributário...")
        texto_consolidado = "\n".join(resumos_volumes)

        instrucao_sistema_mestre = """
        Você é o M.A | JUS IA EXPERIENCE, um Auditor Forense Especialista em Direito Tributário.
        
        >>> TRAVA ABSOLUTA ANTI-ALUCINAÇÃO (TOLERÂNCIA ZERO) <<<
        1. PROIBIDO inventar fatos, CDAs, valores, impostos ou taxas que não estejam EXPLICITAMENTE ESCRITOS nos dados.
        2. VALIDAR APENAS legislação tributária (CTN, CF/88, Leis de Impostos Específicos).
        3. É OBRIGATÓRIO citar as folhas (fls.) de cada achado nos autos.
        4. Se a informação não existir, responda "INFORMAÇÃO NÃO ENCONTRADA".
        
        >>> RIGOR ESTATÍSTICO OAB (PESO DE CORREÇÃO) <<<
        Foque intensamente na estrutura de: Mandado de Segurança (Rigor 8), Agravo (Rigor 7), Apelação (Rigor 7), Repetição de Indébito (Rigor 6).
        """
        
        prompt_mestre = f"DIRECIONAMENTO:\n{fatos}\n\n--- DADOS COLETADOS ---\n{texto_consolidado}\n\nGere o Dossiê Tributário em formato JSON: {{'resumo_estrategico': '...', 'analise_foco_cirurgico': [], 'jurimetria': '...', 'resumo_cliente': '...', 'timeline': [], 'vulnerabilidades_contraparte': [], 'checklist': [], 'base_legal': [], 'jurisprudencia': [], 'doutrina': []}}"

        config_mestre = types.GenerateContentConfig(
            system_instruction=instrucao_sistema_mestre, 
            temperature=0.0, 
            response_mime_type="application/json",
            tools=[{"googleSearch": {}}],
            safety_settings=filtros_seguranca
        )

        response_final = client.models.generate_content(model='gemini-2.0-flash', contents=[prompt_mestre], config=config_mestre)

        if not hasattr(response_final, 'text') or not response_final.text: raise Exception("Recusado pelos filtros.")

        dados_brutos = json.loads(response_final.text.strip(), strict=False)
        dados_sanitizados = aplicar_escudo_no_json(dados_brutos)
        
        callback_sucesso(dados_sanitizados)

    except Exception as e:
        callback_erro(str(e))

def gerar_dossie_word(dados: Dict[str, Any], caminho_salvar: str) -> bool:
    try:
        doc = docx.Document()
        for s in doc.sections: s.top_margin, s.bottom_margin, s.left_margin, s.right_margin = Cm(3), Cm(2), Cm(3), Cm(2)
        doc.add_heading('M.A | DOSSIÊ ESTRATÉGICO TRIBUTÁRIO', level=1)
        
        def add_secao(titulo: str, conteudo: Any):
            if conteudo:
                doc.add_heading(titulo, level=2)
                if isinstance(conteudo, list):
                    for item in conteudo: doc.add_paragraph(str(item), style='List Bullet')
                else:
                    p = doc.add_paragraph(str(conteudo))
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        add_secao('Resumo Estratégico e Probabilidade', dados.get('resumo_estrategico', ''))
        add_secao('Foco Cirúrgico (Auditoria Específica)', dados.get('analise_foco_cirurgico', []))
        add_secao('Jurimetria', dados.get('jurimetria', ''))
        add_secao('Report para o Cliente', dados.get('resumo_cliente', ''))

        if dados.get('timeline'):
            doc.add_heading('Cronologia dos Fatos', level=2)
            for t in dados.get('timeline', []): 
                evento = t.get('evento', '') if isinstance(t, dict) else str(t)
                data = t.get('data', '') if isinstance(t, dict) else ''
                doc.add_paragraph(f"{data}: {evento}" if data else evento, style='List Bullet')

        add_secao('Modo Combate (Nulidade de CDA e Prescrição)', dados.get('vulnerabilidades_contraparte', []))
        add_secao('Providências da Defesa', dados.get('checklist', []))
        add_secao('Base Legal Tributária (CTN/CF)', dados.get('base_legal', []))
        add_secao('Jurisprudência Localizada', dados.get('jurisprudencia', []))
        add_secao('Doutrina', dados.get('doutrina', []))

        doc.save(caminho_salvar)
        return True
    except: return False

janela = tk.Tk()
janela.title("M.A | Auditor Tributário (Data-Driven Rigor)")
janela.geometry("1100x750")
janela.configure(bg="#020617")
dados_analise_atual = {}

style = ttk.Style()
style.theme_use('clam')
style.configure("TNotebook", background="#020617", borderwidth=0)
style.configure("TNotebook.Tab", background="#0f172a", foreground="#94a3b8", padding=[20, 10], font=("Segoe UI", 10, "bold"), borderwidth=0)
# Amarelo/Ouro Tributário
style.map("TNotebook.Tab", background=[("selected", "#ca8a04")], foreground=[("selected", "white")])

def executar_analise():
    fatos = txt_fatos.get(1.0, tk.END).strip()
    caminhos_txt = filedialog.askopenfilenames(filetypes=[("Arquivos de Texto (Volumes)", "*.txt")])
    if not caminhos_txt: return
    caminhos_txt = sorted(list(caminhos_txt))

    btn_executar.config(state=tk.DISABLED, text="⏳ AUDITANDO COM RIGOR OAB TRIBUTÁRIO...", bg="#64748b")
    progress_bar.pack(side=tk.LEFT, padx=10)
    progress_bar['maximum'] = len(caminhos_txt)
    progress_bar['value'] = 0
    
    def atualizar_progresso(atual, total, mensagem):
        def _ui_update():
            lbl_status.config(text=mensagem)
            progress_bar['value'] = atual
        janela.after(0, _ui_update)

    def sucesso(dados):
        def _update_ui_sucesso():
            progress_bar.pack_forget() 
            global dados_analise_atual
            dados_analise_atual = dados
            
            txt_resumo.delete(1.0, tk.END)
            txt_resumo.insert(tk.END, f"=== ESTRATÉGIA ===\n{dados.get('resumo_estrategico','')}\n\n=== JURIMETRIA ===\n{dados.get('jurimetria','')}\n\n=== REPORT CLIENTE ===\n{dados.get('resumo_cliente','')}")
            txt_foco.delete(1.0, tk.END)
            txt_foco.insert(tk.END, "=== ANÁLISE CIRÚRGICA ===\n\n" + "\n\n".join([f"📌 {a}" for a in dados.get('analise_foco_cirurgico', [])]))
            txt_vuln.delete(1.0, tk.END)
            txt_vuln.insert(tk.END, "=== NULIDADES E PRESCRIÇÃO ===\n" + "\n".join([f"- {v}" for v in dados.get('vulnerabilidades_contraparte', [])]) + "\n\n")
            
            timeline_str = ""
            for t in dados.get('timeline', []):
                if isinstance(t, dict): timeline_str += f"[{t.get('data', '')}] {t.get('evento', '')}\n"
                else: timeline_str += f"- {t}\n"
            txt_vuln.insert(tk.END, "=== TIMELINE ===\n" + timeline_str)
            
            txt_legal.delete(1.0, tk.END)
            txt_legal.insert(tk.END, "=== JURISPRUDÊNCIA (SANITIZADA) ===\n" + "\n\n".join([f"- {j}" for j in dados.get('jurisprudencia', [])]) + "\n\n")
            txt_legal.insert(tk.END, "=== BASE LEGAL TRIBUTÁRIA ===\n" + "\n".join([f"- {b}" for b in dados.get('base_legal', [])]) + "\n\n")

            btn_executar.config(state=tk.NORMAL, text="🚀 ANALISAR NOVO LOTE TRIBUTÁRIO", bg="#ca8a04")
            lbl_status.config(text="Status: Auditoria Tributária concluída.", fg="#10b981")
            notebook.select(2)
        janela.after(0, _update_ui_sucesso)

    def erro(msg):
        def _update_ui_erro():
            progress_bar.pack_forget() 
            btn_executar.config(state=tk.NORMAL, text="🚀 TENTAR NOVAMENTE", bg="#ca8a04")
            lbl_status.config(text="Status: Erro na análise.", fg="#ef4444")
            messagebox.showerror("Erro", f"Falha ao auditar:\n{msg}")
        janela.after(0, _update_ui_erro)

    threading.Thread(target=analisar_lote_ia, args=(caminhos_txt, fatos, atualizar_progresso, sucesso, erro), daemon=True).start()

sidebar = tk.Frame(janela, bg="#0f172a", width=300); sidebar.pack(side=tk.LEFT, fill=tk.Y); sidebar.pack_propagate(False)
tk.Label(sidebar, text="M.A | TRIBUTÁRIO", font=("Segoe UI", 18, "bold"), bg="#0f172a", fg="#ca8a04", pady=30).pack()
tk.Label(sidebar, text="Especialista em CTN,\nrepetição e CDAs.\n[Anti-Alucinação CTN]", bg="#0f172a", fg="#cbd5e1", justify="center").pack(pady=10)

main_area = tk.Frame(janela, bg="#020617"); main_area.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True, padx=20, pady=20)
notebook = ttk.Notebook(main_area); notebook.pack(fill=tk.BOTH, expand=True, pady=(0, 15))

tab_fatos = tk.Frame(notebook, bg="#1e293b"); notebook.add(tab_fatos, text=" 1. DIRECIONAMENTO ")
txt_fatos = scrolledtext.ScrolledText(tab_fatos, bg="#1e293b", fg="white", insertbackground="white", borderwidth=0, font=("Segoe UI", 11), padx=15, pady=15)
txt_fatos.pack(fill=tk.BOTH, expand=True); txt_fatos.insert(tk.END, "Fatos e estratégia tributária...")

tab_res = tk.Frame(notebook, bg="#1e293b"); notebook.add(tab_res, text=" 2. ESTRATÉGIA ")
txt_resumo = scrolledtext.ScrolledText(tab_res, bg="#1e293b", fg="#e2e8f0", borderwidth=0, font=("Segoe UI", 11), padx=15, pady=15); txt_resumo.pack(fill=tk.BOTH, expand=True)

tab_foco = tk.Frame(notebook, bg="#1e293b"); notebook.add(tab_foco, text=" 3. FOCO CIRÚRGICO ")
txt_foco = scrolledtext.ScrolledText(tab_foco, bg="#1e293b", fg="#fbbf24", borderwidth=0, font=("Segoe UI", 11), padx=15, pady=15); txt_foco.pack(fill=tk.BOTH, expand=True)

tab_vuln = tk.Frame(notebook, bg="#1e293b"); notebook.add(tab_vuln, text=" 4. NULIDADES / CDA ")
txt_vuln = scrolledtext.ScrolledText(tab_vuln, bg="#1e293b", fg="#60a5fa", borderwidth=0, font=("Segoe UI", 11), padx=15, pady=15); txt_vuln.pack(fill=tk.BOTH, expand=True)

tab_leg = tk.Frame(notebook, bg="#1e293b"); notebook.add(tab_leg, text=" 5. FUNDAMENTAÇÃO ")
txt_legal = scrolledtext.ScrolledText(tab_leg, bg="#1e293b", fg="#34d399", borderwidth=0, font=("Segoe UI", 11), padx=15, pady=15); txt_legal.pack(fill=tk.BOTH, expand=True)

action_frame = tk.Frame(main_area, bg="#020617"); action_frame.pack(fill=tk.X, side=tk.BOTTOM)
btn_executar = tk.Button(action_frame, text="🚀 AUDITAR TRIBUTÁRIO (PADRÃO 1.4.5)", command=executar_analise, bg="#ca8a04", fg="white", font=("Segoe UI", 11, "bold"), pady=12, cursor="hand2")
btn_executar.pack(side=tk.LEFT, fill=tk.X, expand=True)

progress_bar = ttk.Progressbar(action_frame, mode='determinate', length=200)
lbl_status = tk.Label(action_frame, text="Status: Aguardando...", bg="#020617", fg="#64748b", font=("Segoe UI", 10)); lbl_status.pack(side=tk.LEFT, padx=20)

janela.mainloop()
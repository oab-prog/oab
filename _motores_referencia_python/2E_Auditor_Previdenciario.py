import os
import json
import time
import datetime
import threading
import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext, ttk
import docx
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from google import genai
from google.genai import types
import re # O motor do Escudo Anti-Alucinação

API_KEY_LOCAL = "AIzaSyA_riNjroJk8CtkvD2bbFVUw9vidFRzV24" 

# ==========================================
# O ESCUDO DE CÓDIGO (REGEX) - A PORTA DE AÇO
# ==========================================
def aplicar_escudo_no_json(dados):
    # REGEX para capturar siglas previdenciárias essenciais e jurisprudências do TRF/TNU.
    # NB, NIT, DER, APS, TRF, TNU
    padroes_escudo = [
        r'\b(NB|NIT|DER|APS|TRF|TNU)\b', # Siglas solicitadas (Zero Hallucination 2.0)
        r'(?i)\b(?:HC|REsp|RHC|AgRg|AREsp|Apelação|Agravo|AI|Processo n[º°]?\s*)\s*[\d\.\-\/]+\/[A-Z]{2}\b' # Jurisprudências
    ]
    
    padrao_mestre = re.compile('|'.join(padroes_escudo))

    if isinstance(dados, dict):
        return {k: aplicar_escudo_no_json(v) for k, v in dados.items()}
    elif isinstance(dados, list):
        return [aplicar_escudo_no_json(item) for item in dados]
    elif isinstance(dados, str):
        # Cláusula de Veto: Se houver dúvida sobre tempo de contribuição ou carência, responda: INFORMAÇÃO NÃO ENCONTRADA
        if "dúvida" in dados.lower() or "não encontrado" in dados.lower() or "incerto" in dados.lower():
            return "INFORMAÇÃO NÃO ENCONTRADA"
        
        # Aplica o escudo padrão (siglas/jurisprudência)
        texto_escudado = padrao_mestre.sub(lambda m: f"[{m.group(0).upper()}]", dados)
        
        # Regra de Veto: Se encontrar um padrão de valor monetário ou data longa sem contexto de FL, mascara se houver dúvida.
        if re.search(r'\b(\d{1,3}(?:\.\d{3})*,\d{2}|\d{1,2}/\d{1,2}/\d{4})\b', texto_escudado) and not re.search(r'\bfls\.', texto_escudado, re.IGNORECASE):
             return "INFORMAÇÃO NÃO ENCONTRADA"
             
        return texto_escudado
    else:
        return dados

def analisar_lote_ia(caminhos_txt, fatos, callback_progresso, callback_sucesso, callback_erro):
    try:
        client = genai.Client(api_key=API_KEY_LOCAL.strip())
        filtros_seguranca = [
            types.SafetySetting(category="HARM_CATEGORY_HATE_SPEECH", threshold="BLOCK_NONE"),
            types.SafetySetting(category="HARM_CATEGORY_HARASSMENT", threshold="BLOCK_NONE"),
            types.SafetySetting(category="HARM_CATEGORY_SEXUALLY_EXPLICIT", threshold="BLOCK_NONE"),
            types.SafetySetting(category="HARM_CATEGORY_DANGEROUS_CONTENT", threshold="BLOCK_NONE"),
        ]

        resumos_volumes = []
        total_vols = len(caminhos_txt)

        # FASE 1: DETETIVES PREVIDENCIÁRIOS (Análise de CNIS e Laudos)
        for idx, caminho in enumerate(caminhos_txt):
            vol_num = idx + 1
            callback_progresso(vol_num, total_vols, f"Fase 1: Auditando CNIS, PPP e Laudos no Vol {vol_num} de {total_vols}...")
            
            gemini_file = client.files.upload(file=caminho, config={'mime_type': 'text/plain'})
            while True:
                f_info = client.files.get(name=gemini_file.name)
                if "FAILED" in str(f_info.state).upper(): raise Exception(f"Falha no Volume {vol_num}.")
                if "ACTIVE" in str(f_info.state).upper(): break
                time.sleep(2)
            
            conteudo_txt = types.Part.from_uri(file_uri=f_info.uri, mime_type='text/plain')
            
            prompt_detetive = f"DIRECIONAMENTO DA ESTRATÉGIA:\n{fatos}\n\nLeia este volume Previdenciário. Extraia informações precisas sobre: indeferimentos do INSS, divergências no CNIS, laudos periciais (CIDs), PPP (Perfil Profissiográfico Previdenciário), carência e qualidade de segurado. É OBRIGATÓRIO citar as fls. de cada achado. Seja frio e médico-legal. NÃO INVENTE DOENÇAS OU DADOS. Se houver dúvida sobre tempo de contribuição ou carência, responda: INFORMAÇÃO NÃO ENCONTRADA."
            config_detetive = types.GenerateContentConfig(temperature=0.0, safety_settings=filtros_seguranca)
            response_detetive = client.models.generate_content(model='gemini-2.5-flash', contents=[conteudo_txt, prompt_detetive], config=config_detetive)
            
            resumos_volumes.append(f"--- ACHADOS DO VOLUME {vol_num} ---\n{response_detetive.text.strip()}\n")
            client.files.delete(name=gemini_file.name)

        # FASE 2: MESTRE JSON PREVIDENCIÁRIO
        callback_progresso(total_vols, total_vols, "Fase 2: Estruturando Dossiê Jurídico contra o INSS...")
        texto_consolidado = "\n".join(resumos_volumes)

        instrucao_sistema_mestre = """
        Você é o M.A | JUS IA EXPERIENCE, um Auditor Forense Especialista em Direito Previdenciário.
        
        >>> TRAVA ABSOLUTA ANTI-ALUCINAÇÃO (TOLERÂNCIA ZERO) <<<
        1. PROIBIDO inventar doenças (CIDs), peritos, formulários PPP, períodos de contribuição ou requerimentos no INSS que não estejam EXPLICITAMENTE ESCRITOS nos dados.
        2. PROIBIDO inventar números de processos da Justiça Federal (TRF, TNU, STJ).
        3. Se houver dúvida sobre tempo de contribuição ou carência, responda: INFORMAÇÃO NÃO ENCONTRADA.
        
        >>> REGRA DE FORMATAÇÃO JSON <<<
        O retorno DEVE ser EXCLUSIVAMENTE um objeto JSON válido.
        USE ESTA ESTRUTURA:
        {
          "resumo_estrategico": "Diagnóstico do processo contra o INSS citando fls.",
          "analise_foco_cirurgico": ["Ponto previdenciário 1 com fls.", "Ponto 2 com fls."],
          "jurimetria": "Análise da vara ou juizado especial federal.",
          "resumo_cliente": "E-mail explicativo para o segurado/cliente.",
          "timeline": [{"data": "DD/MM/AAAA", "evento": "Descrição (DER, DCB, Perícia) e fls."}],
          "vulnerabilidades_contraparte": ["Falhas do perito, erros de cálculo do INSS ou PPP irregular com fls."],
          "checklist": ["Providência (Ex: Impugnar laudo, pedir complementação de PPP)"],
          "base_legal": ["Lei 8.213/91, Instrução Normativa (IN/PRES) ou Súmula do STJ aplicada"],
          "jurisprudencia": ["Sempre use a tag de omissão ou copie a ementa real fornecida nos autos."],
          "doutrina": ["Tese doutrinária de apoio."]
        }
        """
        
        prompt_mestre = f"DIRECIONAMENTO:\n{fatos}\n\n--- DADOS COLETADOS ---\n{texto_consolidado}\n\nGere o Dossiê Previdenciário em formato JSON estrito."

        config_mestre = types.GenerateContentConfig(
            system_instruction=instrucao_sistema_mestre, 
            temperature=0.0, # <<< ZERO HALLUCINATION 2.0: Determinismo Total >>>
            tools=[{"googleSearch": {}}],
            safety_settings=filtros_seguranca,
            response_mime_type="application/json" # <<< JSON NATIVO >>>
        )

        response_final = client.models.generate_content(model='gemini-2.5-flash', contents=[prompt_mestre], config=config_mestre)

        if not hasattr(response_final, 'text') or not response_final.text: raise Exception("Recusado pelos filtros.")

        # APLICAÇÃO DA PORTA DE AÇO (REGEX) NO JSON BRUTO
        texto_puro = response_final.text.strip()
        
        # JSON Nativo: Remover marcação de Markdown se presente, garantindo que o texto seja puro JSON
        if texto_puro.startswith("```json"): texto_puro = texto_puro[7:]
        elif texto_puro.startswith("```"): texto_puro = texto_puro[3:]
        if texto_puro.endswith("```"): texto_puro = texto_puro[:-3]
        
        # Aplica o escudo sanitizante nos dados já convertidos (no valor da string)
        dados_raw = json.loads(texto_puro.strip(), strict=False)
        dados_escudados = aplicar_escudo_no_json(dados_raw)
        
        callback_sucesso(dados_escudados)

    except Exception as e:
        callback_erro(str(e))

def gerar_dossie_word(dados, caminho_salvar):
    try:
        doc = docx.Document()
        for s in doc.sections: s.top_margin, s.bottom_margin, s.left_margin, s.right_margin = Cm(3), Cm(2), Cm(3), Cm(2)
        doc.add_heading('M.A | DOSSIÊ ESTRATÉGICO PREVIDENCIÁRIO (INSS) - ZERO HALLUCINATION 2.0', level=1)
        
        def add_secao(titulo, conteudo):
            if conteudo:
                doc.add_heading(titulo, level=2)
                p = doc.add_paragraph(conteudo)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        add_secao('Resumo Estratégico e Probabilidade', dados.get('resumo_estrategico', ''))
        if dados.get('analise_foco_cirurgico'):
            doc.add_heading('Foco Cirúrgico (CNIS, Carência, Laudos e Qualidade de Segurado)', level=2)
            for item in dados.get('analise_foco_cirurgico', []): doc.add_paragraph(item, style='List Bullet')
        add_secao('Jurimetria (TRF/TNU)', dados.get('jurimetria', ''))
        add_secao('Report para o Segurado', dados.get('resumo_cliente', ''))

        if dados.get('timeline'):
            doc.add_heading('Cronologia (Requerimentos e Perícias)', level=2)
            for t in dados.get('timeline', []): doc.add_paragraph(f"{t.get('data', '')}: {t.get('evento', '')}", style='List Bullet')
        if dados.get('vulnerabilidades_contraparte'):
            doc.add_heading('Modo Combate (Erros do INSS e Peritos)', level=2)
            for v in dados.get('vulnerabilidades_contraparte', []): doc.add_paragraph(v, style='List Bullet')
        if dados.get('checklist'):
            doc.add_heading('Providências da Defesa', level=2)
            for c in dados.get('checklist', []): doc.add_paragraph(c, style='List Bullet')
        if dados.get('base_legal'):
            doc.add_heading('Base Legal e Normativa', level=2)
            for b in dados.get('base_legal', []): doc.add_paragraph(b, style='List Bullet')
        if dados.get('jurisprudencia'):
            doc.add_heading('Jurisprudência Localizada', level=2)
            for j in dados.get('jurisprudencia', []): doc.add_paragraph(j, style='List Bullet')
        if dados.get('doutrina'):
            doc.add_heading('Doutrina Especializada', level=2)
            for d in dados.get('doutrina', []): doc.add_paragraph(d, style='List Bullet')

        doc.save(caminho_salvar)
        return True
    except: return False

janela = tk.Tk()
janela.title("M.A | Auditor Previdenciário (Blindado)")
janela.geometry("1100x750")
janela.configure(bg="#020617")
dados_analise_atual = {}

style = ttk.Style()
style.theme_use('clam')
style.configure("TNotebook", background="#020617", borderwidth=0)
style.configure("TNotebook.Tab", background="#0f172a", foreground="#94a3b8", padding=[20, 10], font=("Segoe UI", 10, "bold"), borderwidth=0)
# Verde Previdenciário (INSS)
style.map("TNotebook.Tab", background=[("selected", "#15803d")], foreground=[("selected", "white")])

def executar_analise():
    fatos = txt_fatos.get(1.0, tk.END).strip()
    caminhos_txt = filedialog.askopenfilenames(filetypes=[("Arquivos de Texto (Volumes)", "*.txt")])
    if not caminhos_txt: return
    caminhos_txt = sorted(list(caminhos_txt))

    btn_executar.config(state=tk.DISABLED, text="⏳ AUDITANDO VOLUMES COM FILTRO REGEX...", bg="#64748b")
    progress_bar.pack(side=tk.LEFT, padx=10)
    progress_bar['maximum'] = len(caminhos_txt)
    progress_bar['value'] = 0
    
    def atualizar_progresso(atual, total, mensagem):
        def _ui_update():
            lbl_status.config(text=mensagem)
            progress_bar['value'] = atual
        janela.after(0, _ui_update)

    def sucesso(dados):
        def _update_ui_sucesso():
            progress_bar.pack_forget() 
            global dados_analise_atual
            dados_analise_atual = dados
            
            txt_resumo.delete(1.0, tk.END)
            txt_resumo.insert(tk.END, f"=== ESTRATÉGIA ===\n{dados.get('resumo_estrategico','')}\n\n=== JURIMETRIA ===\n{dados.get('jurimetria','')}\n\n=== REPORT CLIENTE ===\n{dados.get('resumo_cliente','')}")
            txt_foco.delete(1.0, tk.END)
            txt_foco.insert(tk.END, "=== ANÁLISE CIRÚRGICA ===\n\n" + "\n\n".join([f"📌 {a}" for a in dados.get('analise_foco_cirurgico', [])]))
            txt_vuln.delete(1.0, tk.END)
            txt_vuln.insert(tk.END, "=== ERROS DO INSS / LAUDOS IMPUGNÁVEIS ===\n" + "\n".join([f"- {v}" for v in dados.get('vulnerabilidades_contraparte', [])]) + "\n\n")
            txt_vuln.insert(tk.END, "=== TIMELINE (DER) ===\n" + "\n".join([f"[{t.get('data', '')}] {t.get('evento', '')}" for t in dados.get('timeline', [])]))
            txt_legal.delete(1.0, tk.END)
            txt_legal.insert(tk.END, "=== JURISPRUDÊNCIA (BLINDADA) ===\n" + "\n\n".join([f"- {j}" for j in dados.get('jurisprudencia', [])]) + "\n\n")
            txt_legal.insert(tk.END, "=== BASE LEGAL PREVIDENCIÁRIA ===\n" + "\n".join([f"- {b}" for b in dados.get('base_legal', [])]) + "\n\n")

            btn_executar.config(state=tk.NORMAL, text="🚀 ANALISAR NOVO LOTE DE VOLUMES", bg="#15803d")
            lbl_status.config(text="Status: Auditoria concluída e blindada com sucesso.", fg="#10b981")
            notebook.select(2)
        janela.after(0, _update_ui_sucesso)

    def erro(msg):
        def _update_ui_erro():
            progress_bar.pack_forget() 
            btn_executar.config(state=tk.NORMAL, text="🚀 TENTAR NOVAMENTE", bg="#15803d")
            lbl_status.config(text="Status: Erro na análise.", fg="#ef4444")
            messagebox.showerror("Erro", f"Falha ao auditar:\n{msg}")
        janela.after(0, _update_ui_erro)

    threading.Thread(target=analisar_lote_ia, args=(caminhos_txt, fatos, atualizar_progresso, sucesso, erro), daemon=True).start()

def salvar_dossie():
    if not dados_analise_atual: return messagebox.showinfo("Aviso", "Faça uma análise primeiro.")
    
    texto_total = json.dumps(dados_analise_atual)
    if "[" in texto_total or "]" in texto_total:
        resposta = messagebox.askyesno("Atenção", "O Dossiê possui marcações de colchetes '[ ]' que indicam informações a preencher. Deseja salvar o Word mesmo assim?")
        if not resposta: return

    caminho = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Document", "*.docx")], initialfile=f"Dossie_Previdenciario_Blindado_MA.docx")
    if caminho:
        gerar_dossie_word(dados_analise_atual, caminho)
        messagebox.showinfo("Sucesso", "Dossiê Previdenciário salvo com sucesso!")

sidebar = tk.Frame(janela, bg="#0f172a", width=300); sidebar.pack(side=tk.LEFT, fill=tk.Y); sidebar.pack_propagate(False)
tk.Label(sidebar, text="M.A | PREVIDENCIÁRIO", font=("Segoe UI", 16, "bold"), bg="#0f172a", fg="#15803d", pady=30).pack()
tk.Label(sidebar, text="Especialista em BPC/LOAS,\naposentadorias e impugnação\nde laudos do INSS.\n[Proteção Regex Ativa]", bg="#0f172a", fg="#cbd5e1", justify="center").pack(pady=10)
tk.Button(sidebar, text="📥 BAIXAR DOSSIÊ (WORD)", command=salvar_dossie, bg="#15803d", fg="white", font=("Segoe UI", 10, "bold"), pady=10).pack(side=tk.BOTTOM, fill="x", padx=20, pady=20)

main_area = tk.Frame(janela, bg="#020617"); main_area.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True, padx=20, pady=20)
notebook = ttk.Notebook(main_area); notebook.pack(fill=tk.BOTH, expand=True, pady=(0, 15))

tab_fatos = tk.Frame(notebook, bg="#1e293b"); notebook.add(tab_fatos, text=" 1. DIRECIONAMENTO ")
txt_fatos = scrolledtext.ScrolledText(tab_fatos, bg="#1e293b", fg="white", insertbackground="white", borderwidth=0, font=("Segoe UI", 11), padx=15, pady=15)
txt_fatos.pack(fill=tk.BOTH, expand=True); txt_fatos.insert(tk.END, "Ex: O INSS negou Aposentadoria Especial alegando EPI Eficaz. Analise o PPP anexado para quebrar a tese do INSS e busque a documentação médica de ruído...")

tab_res = tk.Frame(notebook, bg="#1e293b"); notebook.add(tab_res, text=" 2. ESTRATÉGIA & JURIMETRIA ")
txt_resumo = scrolledtext.ScrolledText(tab_res, bg="#1e293b", fg="#e2e8f0", borderwidth=0, font=("Segoe UI", 11), padx=15, pady=15); txt_resumo.pack(fill=tk.BOTH, expand=True)

tab_foco = tk.Frame(notebook, bg="#1e293b"); notebook.add(tab_foco, text=" 3. FOCO CIRÚRGICO ")
txt_foco = scrolledtext.ScrolledText(tab_foco, bg="#1e293b", fg="#fbbf24", borderwidth=0, font=("Segoe UI", 11), padx=15, pady=15); txt_foco.pack(fill=tk.BOTH, expand=True)

tab_vuln = tk.Frame(notebook, bg="#1e293b"); notebook.add(tab_vuln, text=" 4. ERROS INSS E LAUDOS ")
txt_vuln = scrolledtext.ScrolledText(tab_vuln, bg="#1e293b", fg="#4ade80", borderwidth=0, font=("Segoe UI", 11), padx=15, pady=15); txt_vuln.pack(fill=tk.BOTH, expand=True)

tab_leg = tk.Frame(notebook, bg="#1e293b"); notebook.add(tab_leg, text=" 5. FUNDAMENTAÇÃO ")
txt_legal = scrolledtext.ScrolledText(tab_leg, bg="#1e293b", fg="#34d399", borderwidth=0, font=("Segoe UI", 11), padx=15, pady=15); txt_legal.pack(fill=tk.BOTH, expand=True)

action_frame = tk.Frame(main_area, bg="#020617"); action_frame.pack(fill=tk.X, side=tk.BOTTOM)
btn_executar = tk.Button(action_frame, text="🚀 ANEXAR VOLUMES E AUDITAR (COM TRAVA REGEX)", command=executar_analise, bg="#15803d", fg="white", font=("Segoe UI", 11, "bold"), pady=12, cursor="hand2")
btn_executar.pack(side=tk.LEFT, fill=tk.X, expand=True)

progress_bar = ttk.Progressbar(action_frame, mode='determinate', length=200)
lbl_status = tk.Label(action_frame, text="Status: Aguardando...", bg="#020617", fg="#64748b", font=("Segoe UI", 10)); lbl_status.pack(side=tk.LEFT, padx=20)

janela.mainloop()
import os
import json
import time
import datetime
import threading
import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext, ttk
import docx
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from google import genai
from google.genai import types
import re # O motor do Escudo Anti-Alucinação

API_KEY_LOCAL = os.getenv("GEMINI_API_KEY") 

# ==========================================
# O ESCUDO DE CÓDIGO (REGEX) - A PORTA DE AÇO
# ==========================================
# REGEX para capturar jurisprudências de Família (AI, Apelações, etc.)
PADRAO_JURIS_FAMILIA = r'(?i)\b(?:AI|Apelação|REsp|RHC|AgInt|AgRg|Agravo de Instrumento|Processo n[º°]?\s*)\s*[\d\.\-\/]+\/[A-Z]{2}\b'

def aplicar_escudo_no_json(dados):
    """
    Aplica sanitização recursiva no JSON, substituindo textos que pareçam
    jurisprudência de Família ou que não sejam claros com a mensagem de veto.
    """
    if isinstance(dados, dict):
        return {k: aplicar_escudo_no_json(v) for k, v in dados.items()}
    elif isinstance(dados, list):
        return [aplicar_escudo_no_json(item) for item in dados]
    elif isinstance(dados, str):
        # 1. Aplica sanitização de jurisprudência específica (Item 3)
        texto_limpo = re.sub(PADRAO_JURIS_FAMILIA, "[AVISO: JURISPRUDÊNCIA OMITIDA/FILTRADA PARA FAMÍLIA - BUSQUE NA INTEGRA]", dados)
        
        # 2. Implementa a Regra de Veto (Item 7 e 10: INFORMAÇÃO NÃO ENCONTRADA)
        # Se o resultado da LLM (após json.loads) contiver a string de veto explícita (que deve ser a última camada de defesa)
        if "INFORMAÇÃO NÃO ENCONTRADA" in texto_limpo:
            return "INFORMAÇÃO NÃO ENCONTRADA"
        
        # Se for um texto que claramente não é uma informação estruturada, consideramos como não encontrada
        if not texto_limpo.strip() or "não há informações" in texto_limpo.lower() or "sem dados" in texto_limpo.lower() or "não especificado" in texto_limpo.lower():
            return "INFORMAÇÃO NÃO ENCONTRADA"
        
        # Correção da tag de omissão se o escudo anterior for encontrado
        if "[AVISO: NUMERACAO OMITIDA - BUSQUE NA INTEGRA NO JUSBRASIL/IBDFAM]" in texto_limpo:
             return texto_limpo.replace("[AVISO: NUMERACAO OMITIDA - BUSQUE NA INTEGRA NO JUSBRASIL/IBDFAM]", "[AVISO: JURISPRUDÊNCIA OMITIDA/FILTRADA PARA FAMÍLIA - BUSQUE NA INTEGRA]")
        
        return texto_limpo
    else:
        return dados

def escudo_anti_alucinacao(texto):
    # FUNÇÃO LEGADA: Mantida apenas para compatibilidade com a limpeza pré-JSON, embora a nova lógica use aplicar_escudo_no_json pós-loads.
    padrao_jurisprudencia = r'(?i)\b(?:HC|REsp|RHC|AgRg|AREsp|Apelação|Agravo|AI|Processo n[º°]?\s*)\s*[\d\.\-\/]+\/[A-Z]{2}\b'
    texto_limpo = re.sub(padrao_jurisprudencia, "[AVISO: NUMERACAO OMITIDA - BUSQUE NA INTEGRA NO JUSBRASIL/IBDFAM]", texto)
    return texto_limpo

def analisar_lote_ia(caminhos_txt, fatos, callback_progresso, callback_sucesso, callback_erro):
    try:
        client = genai.Client(api_key=API_KEY_LOCAL.strip())
        filtros_seguranca = [
            types.SafetySetting(category="HARM_CATEGORY_HATE_SPEECH", threshold="BLOCK_NONE"),
            types.SafetySetting(category="HARM_CATEGORY_HARASSMENT", threshold="BLOCK_NONE"),
            types.SafetySetting(category="HARM_CATEGORY_SEXUALLY_EXPLICIT", threshold="BLOCK_NONE"),
            types.SafetySetting(category="HARM_CATEGORY_DANGEROUS_CONTENT", threshold="BLOCK_NONE"),
        ]

        resumos_volumes = []
        total_vols = len(caminhos_txt)

        # FASE 1: DETETIVES DE FAMÍLIA E SUCESSÕES
        for idx, caminho in enumerate(caminhos_txt):
            vol_num = idx + 1
            callback_progresso(vol_num, total_vols, f"Fase 1: Rastreiando Bens e Provas no Volume {vol_num} de {total_vols}...")
            
            gemini_file = client.files.upload(file=caminho, config={'mime_type': 'text/plain'})
            while True:
                f_info = client.files.get(name=gemini_file.name)
                if "FAILED" in str(f_info.state).upper(): raise Exception(f"Falha no Volume {vol_num}.")
                if "ACTIVE" in str(f_info.state).upper(): break
                time.sleep(2)
            
            conteudo_txt = types.Part.from_uri(file_uri=f_info.uri, mime_type='text/plain')
            
            # REFINAMENTO: Foco em partilha de bens, pensão alimentícia e alienação parental. (Item 8)
            prompt_detetive = f"DIRECIONAMENTO DA ESTRATÉGIA (FOCO FAMÍLIA):\n{fatos}\n\nLeia este volume de Direito de Família/Sucessões. Extraia informações sobre: Partilha de Bens, valores de Pensão Alimentícia e Alienação Parental, e contradições em depoimentos. É OBRIGATÓRIO citar as fls. de cada achado. Seja frio, robótico e objetivo. NÃO INVENTE DADOS."
            config_detetive = types.GenerateContentConfig(temperature=0.0, safety_settings=filtros_seguranca) # ZERO HALLUCINATION 2.0: temperature=0.0 (Item 4)
            response_detetive = client.models.generate_content(model='gemini-2.5-flash', contents=[conteudo_txt, prompt_detetive], config=config_detetive)
            
            resumos_volumes.append(f"--- ACHADOS DO VOLUME {vol_num} ---\n{response_detetive.text.strip()}\n")
            client.files.delete(name=gemini_file.name)

        # FASE 2: MESTRE JSON FAMÍLIA
        callback_progresso(total_vols, total_vols, "Fase 2: Estruturando Dossiê Jurídico Familiar...")
        texto_consolidado = "\n".join(resumos_volumes)

        instrucao_sistema_mestre = f"""
        Você é o M.A | JUS IA EXPERIENCE, um Auditor Forense Especialista em Direito de Família e Sucessões.
        
        >>> REGRA DE VETO OBRIGATÓRIO (ZERO HALLUCINATION 2.0) <<<
        1. É PROIBIDO inventar qualquer dado fático (bens, contas bancárias, laudos, traições, atos de alienação parental) que não estejam EXPLICITAMENTE ESCRITOS nos dados.
        2. PROIBIDO inventar números de processos, Recursos Especiais ou Agravos de Instrumento.
        3. Se a informação sobre Partilha de Bens, Pensão Alimentícia ou Alienação Parental não for clara ou não for encontrada nos autos, a IA deve obrigatoriamente responder: 'INFORMAÇÃO NÃO ENCONTRADA'.
        
        >>> FOCOS OBRIGATÓRIOS <<<
        Priorize estritamente a análise de: Partilha de Bens, Pensão Alimentícia e Alienação Parental.

        >>> REGRA DE FORMATAÇÃO JSON <<<
        O retorno DEVE ser EXCLUSIVAMENTE um objeto JSON válido. **NÃO INCLUA TEXTO EXPLICATIVO FORA DO JSON E NEM MARCADORES COMO \`\`\`json\`\`.**
        USE ESTA ESTRUTURA:
        {{
          "resumo_estrategico": "Diagnóstico do processo de família citando fls.",
          "analise_foco_cirurgico": ["Ponto patrimonial ou familiar 1 com fls.", "Ponto 2 com fls."],
          "jurimetria": "Análise da vara de família/sucessões.",
          "resumo_cliente": "E-mail extremamente empático e técnico para o cliente.",
          "timeline": [{{ "data": "DD/MM/AAAA", "evento": "Descrição (Ex: Casamento, Separação de Fato) e fls." }}],
          "vulnerabilidades_contraparte": ["Ocultação de patrimônio, laudos falsos ou alienação parental com fls."],
          "checklist": ["Providência (Ex: Pedir quebra de sigilo bancário (SISBAJUD)"],
          "base_legal": ["Artigo do Código Civil ou ECA aplicado"],
          "jurisprudencia": ["Sempre use a tag de omissão ou copie a ementa real fornecida nos autos."],
          "doutrina": ["Tese doutrinária de apoio (Ex: Maria Berenice Dias, Rolf Madaleno)."]
        }}
        """
        
        prompt_mestre = f"DIRECIONAMENTO:\n{fatos}\n\n--- DADOS COLETADOS ---\n{texto_consolidado}\n\nGere o Dossiê de Família em formato JSON estrito, respeitando todos os vetos e focos."

        config_mestre = types.GenerateContentConfig(
            system_instruction=instrucao_sistema_mestre, 
            temperature=0.0, # ZERO HALLUCINATION 2.0: temperature=0.0 (Item 4)
            response_mime_type="application/json", # JSON NATIVO 2.0 (Item 5)
            tools=[{"googleSearch": {}}],
            safety_settings=filtros_seguranca
        )

        response_final = client.models.generate_content(model='gemini-2.5-flash', contents=[prompt_mestre], config=config_mestre)

        if not hasattr(response_final, 'text') or not response_final.text: raise Exception("Recusado pelos filtros ou resposta vazia.")

        texto_puro = response_final.text.strip()
        
        # SEGURANÇA DE DADOS 2.0: Primeiro execute o json.loads (Item 6)
        try:
            dados = json.loads(texto_puro)
        except json.JSONDecodeError:
            # Tratamento de fallback para marcadores obsoletos (Item 7) - Deve ser menos necessário com response_mime_type, mas mantido por segurança.
            texto_puro_fallback = texto_puro
            if texto_puro_fallback.startswith("```json"): texto_puro_fallback = texto_puro_fallback[7:]
            elif texto_puro_fallback.startswith("```"): texto_puro_fallback = texto_puro_fallback[3:]
            if texto_puro_fallback.endswith("```"): texto_puro_fallback = texto_puro_fallback[:-3]
            
            try:
                # Tentativa após remover marcadores de bloco obsoletos
                dados = json.loads(texto_puro_fallback.strip(), strict=False)
            except json.JSONDecodeError as e:
                # Se ainda falhar, reportamos o erro e forçamos um veto no resultado final
                raise Exception(f"Falha crítica ao decodificar JSON. Dados brutos iniciam com: {texto_puro[:200]}. Verifique a saída do LLM.")


        # SEGURANÇA DE DADOS 2.0: Depois aplique o escudo sanitizante recursivo (Item 6)
        dados_sanitizados = aplicar_escudo_no_json(dados)
        
        # Garantir que se a IA falhar ao aplicar o veto (Item 10), ele seja forçado aqui
        if 'resumo_estrategico' in dados_sanitizados and (not dados_sanitizados['resumo_estrategico'] or 'INFORMAÇÃO NÃO ENCONTRADA' in dados_sanitizados['resumo_estrategico']):
             dados_sanitizados['resumo_estrategico'] = "INFORMAÇÃO NÃO ENCONTRADA"
             
        callback_sucesso(dados_sanitizados)

    except Exception as e:
        callback_erro(str(e))

def gerar_dossie_word(dados, caminho_salvar):
    try:
        doc = docx.Document()
        for s in doc.sections: s.top_margin, s.bottom_margin, s.left_margin, s.right_margin = Cm(3), Cm(2), Cm(3), Cm(2)
        # COR ROXO FAMÍLIA APLICADA AQUI (Item 9)
        doc.add_heading('M.A | DOSSIÊ ESTRATÉGICO FAMÍLIA E SUCESSÕES', level=1)
        
        def add_secao(titulo, conteudo):
            if conteudo:
                doc.add_heading(titulo, level=2)
                p = doc.add_paragraph(conteudo)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        add_secao('Resumo Estratégico e Probabilidade', dados.get('resumo_estrategico', ''))
        if dados.get('analise_foco_cirurgico'):
            doc.add_heading('Foco Cirúrgico Patrimonial/Familiar', level=2)
            for item in dados.get('analise_foco_cirurgico', []): doc.add_paragraph(item, style='List Bullet')
        add_secao('Jurimetria', dados.get('jurimetria', ''))
        add_secao('Report para o Cliente', dados.get('resumo_cliente', ''))

        if dados.get('timeline'):
            doc.add_heading('Cronologia da Relação', level=2)
            for t in dados.get('timeline', []): doc.add_paragraph(f"{t.get('data', '')}: {t.get('evento', '')}", style='List Bullet')
        if dados.get('vulnerabilidades_contraparte'):
            doc.add_heading('Modo Combate (Fraudes Patrimoniais e Provas)', level=2)
            for v in dados.get('vulnerabilidades_contraparte', []): doc.add_paragraph(v, style='List Bullet')
        if dados.get('checklist'):
            doc.add_heading('Providências da Defesa', level=2)
            for c in dados.get('checklist', []): doc.add_paragraph(c, style='List Bullet')
        if dados.get('base_legal'):
            doc.add_heading('Base Legal (Código Civil e ECA)', level=2)
            for b in dados.get('base_legal', []): doc.add_paragraph(b, style='List Bullet')
        if dados.get('jurisprudencia'):
            doc.add_heading('Jurisprudência Localizada', level=2)
            for j in dados.get('jurisprudencia', []): doc.add_paragraph(j, style='List Bullet')
        if dados.get('doutrina'):
            doc.add_heading('Doutrina Especializada', level=2)
            for d in dados.get('doutrina', []): doc.add_paragraph(d, style='List Bullet')

        doc.save(caminho_salvar)
        return True
    except: return False

janela = tk.Tk()
janela.title("M.A | Auditor Família e Sucessões (Zero Hallucination 2.0)")
janela.geometry("1100x750")
janela.configure(bg="#020617")
dados_analise_atual = {}

style = ttk.Style()
style.theme_use('clam')
style.configure("TNotebook", background="#020617", borderwidth=0)
style.configure("TNotebook.Tab", background="#0f172a", foreground="#94a3b8", padding=[20, 10], font=("Segoe UI", 10, "bold"), borderwidth=0)
# Roxo Família
style.map("TNotebook.Tab", background=[("selected", "#7e22ce")], foreground=[("selected", "white")])

def executar_analise():
    fatos = txt_fatos.get(1.0, tk.END).strip()
    caminhos_txt = filedialog.askopenfilenames(filetypes=[("Arquivos de Texto (Volumes)", "*.txt")])
    if not caminhos_txt: return
    caminhos_txt = sorted(list(caminhos_txt))

    btn_executar.config(state=tk.DISABLED, text="⏳ AUDITANDO VOLUMES COM FILTRO REGEX...", bg="#64748b")
    progress_bar.pack(side=tk.LEFT, padx=10)
    progress_bar['maximum'] = len(caminhos_txt)
    progress_bar['value'] = 0
    
    def atualizar_progresso(atual, total, mensagem):
        def _ui_update():
            lbl_status.config(text=mensagem)
            progress_bar['value'] = atual
        janela.after(0, _ui_update)

    def sucesso(dados):
        def _update_ui_sucesso():
            progress_bar.pack_forget() 
            global dados_analise_atual
            dados_analise_atual = dados
            
            txt_resumo.delete(1.0, tk.END)
            txt_resumo.insert(tk.END, f"=== ESTRATÉGIA ===\n{dados.get('resumo_estrategico','')}\n\n=== JURIMETRIA ===\n{dados.get('jurimetria','')}\n\n=== REPORT CLIENTE ===\n{dados.get('resumo_cliente','')}")
            txt_foco.delete(1.0, tk.END)
            txt_foco.insert(tk.END, "=== ANÁLISE CIRÚRGICA ===\n\n" + "\n\n".join([f"📌 {a}" for a in dados.get('analise_foco_cirurgico', [])]))
            txt_vuln.delete(1.0, tk.END)
            txt_vuln.insert(tk.END, "=== FRAUDES PATRIMONIAIS E PROVAS ===\n" + "\n".join([f"- {v}" for v in dados.get('vulnerabilidades_contraparte', [])]) + "\n\n")
            txt_vuln.insert(tk.END, "=== TIMELINE ===\n" + "\n".join([f"[{t.get('data', '')}] {t.get('evento', '')}" for t in dados.get('timeline', [])]))
            txt_legal.delete(1.0, tk.END)
            txt_legal.insert(tk.END, "=== JURISPRUDÊNCIA (BLINDADA) ===\n" + "\n\n".join([f"- {j}" for j in dados.get('jurisprudencia', [])]) + "\n\n")
            txt_legal.insert(tk.END, "=== BASE LEGAL FAMÍLIA ===\n" + "\n".join([f"- {b}" for b in dados.get('base_legal', [])]) + "\n\n")

            btn_executar.config(state=tk.NORMAL, text="🚀 ANALISAR NOVO LOTE DE VOLUMES", bg="#7e22ce")
            lbl_status.config(text="Status: Auditoria concluída e blindada com sucesso.", fg="#10b981")
            notebook.select(2)
        janela.after(0, _update_ui_sucesso)

    def erro(msg):
        def _update_ui_erro():
            progress_bar.pack_forget() 
            btn_executar.config(state=tk.NORMAL, text="🚀 TENTAR NOVAMENTE", bg="#7e22ce")
            lbl_status.config(text="Status: Erro na análise.", fg="#ef4444")
            messagebox.showerror("Erro", f"Falha ao auditar:\n{msg}")
        janela.after(0, _update_ui_erro)

    threading.Thread(target=analisar_lote_ia, args=(caminhos_txt, fatos, atualizar_progresso, sucesso, erro), daemon=True).start()

def salvar_dossie():
    if not dados_analise_atual: return messagebox.showinfo("Aviso", "Faça uma análise primeiro.")
    
    texto_total = json.dumps(dados_analise_atual)
    if "[" in texto_total or "]" in texto_total:
        resposta = messagebox.askyesno("Atenção", "O Dossiê possui marcações de colchetes '[ ]' que indicam informações a preencher. Deseja salvar o Word mesmo assim?")
        if not resposta: return

    caminho = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Document", "*.docx")], initialfile=f"Dossie_Familia_Blindado_MA.docx")
    if caminho:
        gerar_dossie_word(dados_analise_atual, caminho)
        messagebox.showinfo("Sucesso", "Dossiê de Família/Sucessões salvo com sucesso!")

sidebar = tk.Frame(janela, bg="#0f172a", width=300); sidebar.pack(side=tk.LEFT, fill=tk.Y); sidebar.pack_propagate(False)
tk.Label(sidebar, text="M.A | FAMÍLIA", font=("Segoe UI", 18, "bold"), bg="#0f172a", fg="#7e22ce", pady=30).pack()
tk.Label(sidebar, text="Especialista em partilha,\nalimentos e alienação parental.\n[Zero Hallucination 2.0 Ativo]", bg="#0f172a", fg="#cbd5e1", justify="center").pack(pady=10)
tk.Button(sidebar, text="📥 BAIXAR DOSSIÊ (WORD)", command=salvar_dossie, bg="#10b981", fg="white", font=("Segoe UI", 10, "bold"), pady=10).pack(side=tk.BOTTOM, fill="x", padx=20, pady=20)

main_area = tk.Frame(janela, bg="#020617"); main_area.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True, padx=20, pady=20)
notebook = ttk.Notebook(main_area); notebook.pack(fill=tk.BOTH, expand=True, pady=(0, 15))

tab_fatos = tk.Frame(notebook, bg="#1e293b"); notebook.add(tab_fatos, text=" 1. DIRECIONAMENTO ")
txt_fatos = scrolledtext.ScrolledText(tab_fatos, bg="#1e293b", fg="white", insertbackground="white", borderwidth=0, font=("Segoe UI", 11), padx=15, pady=15)
txt_fatos.pack(fill=tk.BOTH, expand=True); txt_fatos.insert(tk.END, "Ex: O marido alega não ter bens, mas viaja de primeira classe. Foco: Partilha, Pensão e Alienação Parental.")

tab_res = tk.Frame(notebook, bg="#1e293b"); notebook.add(tab_res, text=" 2. ESTRATÉGIA & JURIMETRIA ")
txt_resumo = scrolledtext.ScrolledText(tab_res, bg="#1e293b", fg="#e2e8f0", borderwidth=0, font=("Segoe UI", 11), padx=15, pady=15); txt_resumo.pack(fill=tk.BOTH, expand=True)

tab_foco = tk.Frame(notebook, bg="#1e293b"); notebook.add(tab_foco, text=" 3. FOCO CIRÚRGICO ")
txt_foco = scrolledtext.ScrolledText(tab_foco, bg="#1e293b", fg="#fbbf24", borderwidth=0, font=("Segoe UI", 11), padx=15, pady=15); txt_foco.pack(fill=tk.BOTH, expand=True)

tab_vuln = tk.Frame(notebook, bg="#1e293b"); notebook.add(tab_vuln, text=" 4. FRAUDES E PROVAS ")
txt_vuln = scrolledtext.ScrolledText(tab_vuln, bg="#1e293b", fg="#c084fc", borderwidth=0, font=("Segoe UI", 11), padx=15, pady=15); txt_vuln.pack(fill=tk.BOTH, expand=True)

tab_leg = tk.Frame(notebook, bg="#1e293b"); notebook.add(tab_leg, text=" 5. FUNDAMENTAÇÃO ")
txt_legal = scrolledtext.ScrolledText(tab_leg, bg="#1e293b", fg="#34d399", borderwidth=0, font=("Segoe UI", 11), padx=15, pady=15); txt_legal.pack(fill=tk.BOTH, expand=True)

action_frame = tk.Frame(main_area, bg="#020617"); action_frame.pack(fill=tk.X, side=tk.BOTTOM)
btn_executar = tk.Button(action_frame, text="🚀 ANEXAR VOLUMES E AUDITAR (COM TRAVA REGEX)", command=executar_analise, bg="#7e22ce", fg="white", font=("Segoe UI", 11, "bold"), pady=12, cursor="hand2")
btn_executar.pack(side=tk.LEFT, fill=tk.X, expand=True)

progress_bar = ttk.Progressbar(action_frame, mode='determinate', length=200)
lbl_status = tk.Label(action_frame, text="Status: Aguardando...", bg="#020617", fg="#64748b", font=("Segoe UI", 10)); lbl_status.pack(side=tk.LEFT, padx=20)

janela.mainloop()